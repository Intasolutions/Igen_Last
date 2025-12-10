# analytics/services/exports.py
import os
from io import BytesIO
from typing import Iterable, List
from datetime import datetime
from decimal import Decimal

from django.utils import timezone  # use Django timezone
from django.conf import settings

# ---------- Excel ----------
try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
except Exception:
    Workbook = None
    Font = PatternFill = Alignment = Border = Side = None
    get_column_letter = None

# ---------- PDF ----------
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle
except Exception:
    A4 = None
    canvas = None
    cm = None
    colors = None
    Table = None
    TableStyle = None

# ---------- DOCX (Word) ----------
try:
    from docx import Document  # pip install python-docx
    from docx.shared import Pt
except Exception:
    Document = None
    Pt = None


# ------------------------ Helpers ------------------------


def _find_default_logo_path() -> str:
    """
    Try to locate a good default logo on disk.

    Priority:
      1) igen-frontend/public/logo/logo.jpeg   (blue square logo)
      2) igen-frontend/public/logo/igenprpt.jpeg
      3) igen-frontend/public/logo/igen.png
      4) igen-frontend/public/logo/igen.jpeg
      5) assets/igen-logo.png (legacy fallback)
    """
    base = getattr(settings, "BASE_DIR", None)
    if not base:
        return ""

    candidates = [
        ("igen-frontend", "public", "logo", "logo.jpeg"),
        ("igen-frontend", "public", "logo", "igenprpt.jpeg"),
        ("igen-frontend", "public", "logo", "igen.png"),
        ("igen-frontend", "public", "logo", "igen.jpeg"),
        ("assets", "igen-logo.png"),
    ]
    for parts in candidates:
        p = os.path.join(base, *parts)
        if os.path.exists(p):
            return p
    return ""


DEFAULT_LOGO_PATH = _find_default_logo_path()
LOGO_PATH = getattr(settings, "IGEN_PDF_LOGO", DEFAULT_LOGO_PATH)

_MONEY_HEADERS = {
    "credit",
    "debit",
    "balance",
    "inflows",
    "outflows",
    "net",
    "margin",
    "igen_sc_this_month",
    "expected_rent_this_month",
}


def _looks_number(x) -> bool:
    return isinstance(x, (int, float, Decimal))


def _money_col(header: str) -> bool:
    if not header:
        return False
    h = str(header).strip().lower()
    return h in _MONEY_HEADERS


# ------------------------ Excel Export ------------------------


def export_excel(headers: List[str], rows: Iterable[Iterable]):
    """
    Create a styled XLSX workbook and return it as BytesIO.
    Signature kept simple to match existing calls.
    """
    if Workbook is None:
        raise RuntimeError("openpyxl is required (pip install openpyxl)")

    wb = Workbook()
    ws = wb.active
    ws.title = "Report"

    # Header row
    ws.append(list(headers or []))

    # Data rows
    for r in rows:
        ws.append(list(r))

    # Styles if available
    if Font and PatternFill and Alignment and Border and Side and get_column_letter:
        # Header style
        header_font = Font(bold=True)
        header_fill = PatternFill("solid", fgColor="E5E7EB")  # light gray
        thin = Side(style="thin", color="CCCCCC")
        header_border = Border(bottom=thin)

        for col_idx, h in enumerate(headers, start=1):
            cell = ws.cell(row=1, column=col_idx)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(vertical="center")
            cell.border = header_border

        # Freeze header, enable filters
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions

        # Autosize columns & number formats
        max_len = {
            i: len(str(h)) if h is not None else 0
            for i, h in enumerate(headers, start=1)
        }
        money_cols = set()

        for row in ws.iter_rows(min_row=2, values_only=True):
            for i, val in enumerate(row, start=1):
                l = len(str(val)) if val is not None else 0
                if l > max_len.get(i, 0):
                    max_len[i] = l
                # Track potential money columns by header
                if _looks_number(val) and _money_col(
                    headers[i - 1] if i - 1 < len(headers) else ""
                ):
                    money_cols.add(i)

        for i, length in max_len.items():
            col_letter = get_column_letter(i)
            # width heuristic: char count + padding, bounded
            ws.column_dimensions[col_letter].width = min(max(length + 2, 10), 50)

        # Apply number format to "money" columns
        for row_idx in range(2, ws.max_row + 1):
            for col_idx in money_cols:
                ws.cell(row=row_idx, column=col_idx).number_format = "#,##0.00"

        # Row heights (optional subtle)
        ws.row_dimensions[1].height = 18

    # Serialize
    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio


# ------------------------ PDF Export ------------------------


def export_simple_pdf(title: str, headers: List[str], rows: List[List[str]]):
    """
    Styled PDF table for owner / analytics reports:
      - Logo + centered title at top
      - Proper grid table with coloured header
      - Footer with company contact info in a small table
      - Page number + generated timestamp
      - Multi-line cells (no text cut off)

    If reportlab is missing, falls back to a simple text output.
    """
    bio = BytesIO()

    # Text fallback when reportlab is unavailable
    if not (canvas and A4 and cm):
        bio.write((title + "\n").encode())
        bio.write((" | ".join(map(str, headers)) + "\n").encode())
        for r in rows:
            bio.write((" | ".join(map(str, r)) + "\n").encode())
        bio.seek(0)
        return bio

    c = canvas.Canvas(bio, pagesize=A4)
    width, height = A4

    # Margins
    left = 1.8 * cm
    right = 1.8 * cm
    top = 2.5 * cm
    bottom = 2.2 * cm

    BODY_FONT = "Helvetica"
    BODY_SIZE = 9
    LINE_HEIGHT = 11  # distance between wrapped lines

    # ------------ wrapping helper ------------

    def wrap_text(text: str, max_width: float) -> List[str]:
        """
        Wrap text into multiple lines that fit within max_width
        using BODY_FONT/BODY_SIZE.
        """
        s = str(text or "")
        if not s:
            return [""]

        words = s.split()
        lines: List[str] = []
        current = ""

        for word in words:
            tentative = (current + " " + word).strip()
            if c.stringWidth(tentative, BODY_FONT, BODY_SIZE) <= max_width:
                current = tentative
            else:
                if current:
                    lines.append(current)
                # if single long word, split by chars
                if c.stringWidth(word, BODY_FONT, BODY_SIZE) <= max_width:
                    current = word
                else:
                    chunk = ""
                    for ch in word:
                        if c.stringWidth(chunk + ch, BODY_FONT, BODY_SIZE) <= max_width:
                            chunk += ch
                        else:
                            lines.append(chunk)
                            chunk = ch
                    current = chunk
        if current:
            lines.append(current)
        return lines or [""]

    # ------------ title + logo ------------

    def draw_title() -> float:
        y = height - top

        # draw logo if file exists
        if LOGO_PATH and os.path.exists(LOGO_PATH):
            try:
                logo_width = 3.5 * cm
                logo_height = 2.0 * cm
                c.drawImage(
                    str(LOGO_PATH),
                    (width - logo_width) / 2.0,
                    y - logo_height,
                    width=logo_width,
                    height=logo_height,
                    preserveAspectRatio=True,  # keep original colours
                )
                y -= logo_height + 0.4 * cm
            except Exception:
                # fail silently if image can't be loaded
                pass

        c.setFont("Helvetica-Bold", 13)
        c.drawCentredString(width / 2.0, y, title)
        # space below title where table can start from
        return y - 1.0 * cm

    # ------------ footer with contact / page ------------

    def draw_footer(page_num: int):
        now = (
            timezone.localtime()
            if hasattr(timezone, "localtime")
            else datetime.now()
        )
        ts = now.strftime("%Y-%m-%d %H:%M")

        if Table and TableStyle:
            # company contact table (2 columns)
            contact_data = [
                ["Portico Down Town, Seaport Airport Road, Kakkanad, Kochi", ""],
                ["Email:", "info@igenproperties.in"],
                ["Phone:", "+91 6282706378, +91 6282796060"],
                ["Website:", "www.igenproperties.in"],
            ]

            table_width = width - left - right
            # first column ~3.2cm, second column rest
            col_widths = [3.2 * cm, table_width - 3.2 * cm]

            t = Table(contact_data, colWidths=col_widths)
            t.setStyle(
                TableStyle(
                    [
                        ("SPAN", (0, 0), (-1, 0)),  # first row spans 2 columns
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE", (0, 0), (-1, -1), 8),
                        ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
                        ("LEFTPADDING", (0, 0), (-1, -1), 2),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                        ("TOPPADDING", (0, 0), (-1, -1), 1),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
                        ("FONTNAME", (0, 1), (0, -1), "Helvetica-Bold"),
                    ]
                )
            )
            # place contact table slightly above bottom margin
            y_table = bottom - 0.6 * cm
            t.wrapOn(c, table_width, 3 * cm)
            t.drawOn(c, left, y_table)

            y_text = y_table - 0.35 * cm
        else:
            # very simple footer if platypus is missing
            y_text = bottom - 0.3 * cm
            c.setFont("Helvetica", 8)
            c.drawString(
                left,
                y_text + 24,
                "Portico Down Town, Seaport Airport Road, Kakkanad, Kochi",
            )
            c.drawString(left, y_text + 14, "Email:  info@igenproperties.in")
            c.drawString(
                left,
                y_text + 4,
                "Phone:  +91 6282706378, +91 6282796060  Website: www.igenproperties.in",
            )

        # generated + page number line
        c.setFont("Helvetica", 8)
        c.setFillColor(colors.grey)
        c.drawString(left, y_text, f"Generated {ts}")
        c.drawRightString(width - right, y_text, f"Page {page_num}")
        c.setFillColor(colors.black)

    # ------------ table helpers ------------

    n_cols = max(len(headers), max((len(r) for r in rows), default=0))
    usable_width = width - left - right

    # Basic equal widths with wider "Remarks" column if present
    base_col_width = usable_width / max(n_cols, 1)
    col_widths = [base_col_width] * n_cols
    wider_idx = None
    for i, h in enumerate(headers):
        if str(h).strip().lower() in {"remarks", "description", "narration"}:
            wider_idx = i
            break
    if wider_idx is not None and n_cols > 1:
        # make remarks column ~2.4x others
        total_units = n_cols - 1 + 2.4
        unit_w = usable_width / total_units
        for i in range(n_cols):
            col_widths[i] = unit_w * (2.4 if i == wider_idx else 1.0)

    HEADER_BG = colors.HexColor("#fbbf24")  # warm yellow
    ROW_ALT = colors.HexColor("#f9fafb")    # very light grey
    HEADER_HEIGHT = 16

    def draw_header(y_top: float) -> float:
        """Draw header row as coloured band with grid."""
        c.setFont("Helvetica-Bold", 9)
        x = left
        for i in range(n_cols):
            txt = str(headers[i]) if i < len(headers) else f"Col {i+1}"
            w = col_widths[i]

            # cell background
            c.setFillColor(HEADER_BG)
            c.rect(x, y_top - HEADER_HEIGHT, w, HEADER_HEIGHT, stroke=1, fill=1)

            # text
            c.setFillColor(colors.black)
            c.drawString(x + 3, y_top - 11, txt)

            x += w
        return y_top - HEADER_HEIGHT

    def compute_row_layout(row: List[str]):
        """
        For a given logical row, return (cell_lines, row_height)
        where cell_lines is a list[list[str]] (wrapped lines for each cell).
        """
        cell_lines: List[List[str]] = []
        max_lines = 1
        for i in range(n_cols):
            val = row[i] if i < len(row) else ""
            w = col_widths[i] - 6  # padding
            lines = wrap_text(val, max_width=w)
            cell_lines.append(lines)
            if len(lines) > max_lines:
                max_lines = len(lines)
        row_height = max_lines * LINE_HEIGHT + 4  # padding
        return cell_lines, row_height

    def draw_row(y_top: float, cell_lines: List[List[str]], idx: int, row_height: float) -> float:
        """Draw single data row with grid + optional zebra stripes."""
        c.setFont(BODY_FONT, BODY_SIZE)
        x = left
        bg = ROW_ALT if idx % 2 == 1 else colors.white

        for col in range(n_cols):
            w = col_widths[col]
            lines = cell_lines[col]

            # background rect
            c.setFillColor(bg)
            c.rect(x, y_top - row_height, w, row_height, stroke=1, fill=1)

            # text lines
            c.setFillColor(colors.black)
            text_y = y_top - 3 - BODY_SIZE
            for line in lines:
                c.drawString(x + 3, text_y, line)
                text_y -= LINE_HEIGHT

            x += w

        return y_top - row_height

    # ------------ pagination + drawing ------------

    current_y = draw_title()
    current_y -= 0.3 * cm  # gap before header
    page = 1

    current_y = draw_header(current_y)

    for idx, r in enumerate(rows):
        # compute layout to know required height
        cell_lines, row_height = compute_row_layout(r)

        # create a "safe" area above footer
        if current_y - row_height < bottom + 3.0 * cm:  # time for next page
            draw_footer(page)
            c.showPage()
            page += 1
            current_y = draw_title()
            current_y -= 0.3 * cm
            current_y = draw_header(current_y)

        current_y = draw_row(current_y, cell_lines, idx, row_height)

    draw_footer(page)
    c.save()
    bio.seek(0)
    return bio


# ------------------------ DOCX Export ------------------------


def export_simple_docx(title: str, headers: List[str], rows: List[List[str]]) -> BytesIO:
    """
    Minimal Word table export.
    Returns a BytesIO stream of a .docx file.
    Falls back to a plain-text .docx if python-docx is unavailable.
    """
    bio = BytesIO()

    if Document is None:
        # Fallback: write a simple text-based content
        bio.write((title + "\n").encode())
        bio.write((" | ".join(map(str, headers)) + "\n").encode())
        for r in rows:
            bio.write((" | ".join(map(str, r)) + "\n").encode())
        bio.seek(0)
        return bio

    doc = Document()
    doc.add_heading(title, level=1)

    # Table with header row
    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(str(h))
        if Pt:
            run.font.bold = True
            run.font.size = Pt(10)

    # Data rows
    for r in rows:
        cells = table.add_row().cells
        for i in range(len(headers)):
            cells[i].text = str(r[i] if i < len(r) else "")

    doc.save(bio)
    bio.seek(0)
    return bio
