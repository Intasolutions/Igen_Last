# analytics/views.py
from properties.models import Property
from datetime import date, timedelta
from decimal import Decimal
from io import BytesIO  # ← for DOCX response

from django.apps import apps
from django.db.models import Sum, Q
from django.http import HttpResponse
from django.utils.dateparse import parse_date
from rest_framework import permissions
from rest_framework.response import Response
from rest_framework.views import APIView

# Optional Word export (python-docx)
try:
    from docx import Document  # pip install python-docx
    from docx.shared import Pt
    _DOCX_OK = True
except Exception:
    Document = None
    Pt = None
    _DOCX_OK = False

from .serializers import (
    EntityStatementRowSerializer,
    EntityBalanceSerializer,
    OwnerRentalSummarySerializer,
    OwnerRentalRowSerializer,
    ProjectProfitRowSerializer,
)
from .services.exports import export_excel, export_simple_pdf
from .services.ledger import unified_ledger, running_balance, opening_balance_until


# --------------------------- helpers ---------------------------

def _month_range(yyyy_mm: str):
    """
    Return (start, end) where end is the **inclusive** last day of the month.
    """
    y, m = yyyy_mm.split("-")
    y = int(y)
    m = int(m)
    start = date(y, m, 1)
    next_first = date(y + (m == 12), 1 if m == 12 else m + 1, 1)
    end = next_first - timedelta(days=1)
    return start, end


def _dim_value(row, dim: str):
    """
    Read a dimension out of a unified_ledger row.
    Special-case 'date' to use the value_date.
    """
    if dim == "date":
        return row.get("value_date")
    return row.get(dim)


def _format_period(d: date, granularity: str) -> str:
    """
    Map a date to day/month/quarter/year label strings for pivot 'date' dim.
    """
    if not d:
        return "—"
    g = (granularity or "day").lower()
    if g == "day":
        return d.isoformat()
    if g == "month":
        return f"{d.year:04d}-{d.month:02d}"
    if g == "quarter":
        q = (d.month - 1) // 3 + 1
        return f"{d.year:04d}-Q{q}"
    if g == "year":
        return f"{d.year:04d}"
    return d.isoformat()


# --------------------------- Quick health/debug ---------------------------

class AnalyticsHealthView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        # very wide range just to see *something*
        rows = unified_ledger(request.user, from_date=date(2000, 1, 1), to_date=date(2100, 1, 1))
        return Response({"ok": True, "row_count": len(rows)})


class AnalyticsDataProbeView(APIView):
    """
    Debug endpoint: shows discovered models & counts for cash_ledger, tx_classify, bank_uploads.
    """
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        def model_info(app_label):
            out = []
            try:
                cfg = apps.get_app_config(app_label)
            except Exception:
                return {"error": "app not installed"}
            for M in cfg.get_models():
                try:
                    fields = [f.name for f in M._meta.get_fields()]
                    # pick a likely date field
                    date_field = next((d for d in ["value_date", "transaction_date", "date", "posting_date", "book_date"] if d in fields), None)
                    cnt = M.objects.count()
                    sample = []
                    if cnt:
                        sample = list(M.objects.values_list(date_field or "id")[:3])
                    out.append({
                        "model": M.__name__,
                        "count": cnt,
                        "date_field": date_field,
                        "has_credit": "credit" in fields,
                        "has_debit": "debit" in fields,
                        "has_amount": any(a in fields for a in ["amount", "value", "deposit", "withdrawal"]),
                        "fields": fields[:30],
                        "sample": sample,
                    })
                except Exception as e:
                    out.append({"model": M.__name__, "error": str(e)})
            return out

        payload = {
            "cash_ledger": model_info("cash_ledger"),
            "tx_classify": model_info("tx_classify"),
            "bank_uploads": model_info("bank_uploads"),
        }
        return Response(payload)


# --------------------------- NEW: Entity quick search ---------------------------

class EntityQuickSearchView(APIView):
    """
    Read-only helper for the UI to discover entity IDs safely.
    GET params:
      - q: search text (matches name icontains or numeric id)
      - limit: optional (default 20, max 50)
    Response: [{id, name}]
    """
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        q = (request.GET.get("q") or "").strip()
        try:
            limit = min(max(int(request.GET.get("limit", "20")), 1), 50)
        except Exception:
            limit = 20

        try:
            Entity = apps.get_model("entities", "Entity")
        except Exception:
            return Response([], status=200)

        qs = Entity.objects.all().order_by("name")
        if q:
            cond = Q(name__icontains=q)
            if q.isdigit():
                cond |= Q(id=int(q))
            qs = qs.filter(cond)

        rows = list(qs.values("id", "name")[:limit])
        return Response(rows)


# --------------------------- Report 1: Entity-wise Monthly Statement ---------------------------

class EntityStatementView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        entity_id = request.GET.get("entity_id")
        month = request.GET.get("month")  # YYYY-MM
        if not (entity_id and month):
            return Response({"detail": "entity_id & month (YYYY-MM) required"}, status=400)

        start, end = _month_range(month)
        # Opening balance is strictly before start (exclusive)
        obal = opening_balance_until(request.user, start, entity_id=int(entity_id))
        base_rows = unified_ledger(request.user, from_date=start, to_date=end, entity_id=int(entity_id))
        rows = running_balance(base_rows, opening_balance=obal)

        data = [{
            "value_date": r["value_date"],
            "txn_type": r.get("txn_type"),
            "credit": r.get("credit", Decimal("0")),
            "debit": r.get("debit", Decimal("0")),
            "balance": r.get("balance", Decimal("0")),
            "remarks": r.get("remarks") or "",
        } for r in rows]
        return Response(EntityStatementRowSerializer(data, many=True).data)


class EntityStatementPDFView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        entity_id = request.GET.get("entity_id")
        month = request.GET.get("month")
        if not (entity_id and month):
            return Response({"detail": "entity_id & month (YYYY-MM) required"}, status=400)

        start, end = _month_range(month)
        obal = opening_balance_until(request.user, start, entity_id=int(entity_id))
        base_rows = unified_ledger(request.user, from_date=start, to_date=end, entity_id=int(entity_id))
        rows = running_balance(base_rows, opening_balance=obal)

        headers = ["Date", "Type", "Credit", "Debit", "Balance", "Remarks"]
        table = [[
            r["value_date"].strftime("%Y-%m-%d"),
            r.get("txn_type") or "",
            str(r.get("credit") or 0),
            str(r.get("debit") or 0),
            str(r.get("balance") or 0),
            (r.get("remarks") or "")[:64],
        ] for r in rows]

        pdf = export_simple_pdf(f"Entity {entity_id} - Statement {month}", headers, table)
        resp = HttpResponse(pdf.read(), content_type="application/pdf")
        resp["Content-Disposition"] = f'attachment; filename="entity_{entity_id}_{month}_statement.pdf"'
        return resp


class EntityStatementDOCXView(APIView):
    """
    Word export for Entity Statement (Report 1).
    Requires python-docx. Returns 400 if params missing, 500 if python-docx unavailable.
    """
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        if not _DOCX_OK:
            return Response({"detail": "Word export unavailable: install python-docx"}, status=500)

        entity_id = request.GET.get("entity_id")
        month = request.GET.get("month")
        if not (entity_id and month):
            return Response({"detail": "entity_id & month (YYYY-MM) required"}, status=400)

        start, end = _month_range(month)
        obal = opening_balance_until(request.user, start, entity_id=int(entity_id))
        base_rows = unified_ledger(request.user, from_date=start, to_date=end, entity_id=int(entity_id))
        rows = running_balance(base_rows, opening_balance=obal)

        doc = Document()
        doc.add_heading(f"Entity {entity_id} - Monthly Statement", level=1)
        p = doc.add_paragraph()
        p.add_run(f"Period: {month}").italic = True

        # Table: headers + rows
        headers = ["Date", "Type", "Credit", "Debit", "Balance", "Remarks"]
        table = doc.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            run = hdr_cells[i].paragraphs[0].add_run(h)
            if Pt:
                run.font.bold = True
                run.font.size = Pt(10)

        for r in rows:
            row_cells = table.add_row().cells
            row_cells[0].text = r["value_date"].strftime("%Y-%m-%d")
            row_cells[1].text = str(r.get("txn_type") or "")
            row_cells[2].text = str(r.get("credit") or 0)
            row_cells[3].text = str(r.get("debit") or 0)
            row_cells[4].text = str(r.get("balance") or 0)
            row_cells[5].text = (r.get("remarks") or "")[:256]

        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        resp = HttpResponse(
            bio.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        resp["Content-Disposition"] = f'attachment; filename="entity_{entity_id}_{month}_statement.docx"'
        return resp


# --------------------------- Report 2: Maintenance & Interior (YTD) ---------------------------

class MIExpensesSummaryView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        f = parse_date(request.GET.get("from")) or date(date.today().year, 1, 1)
        t = parse_date(request.GET.get("to")) or date.today()
        rows = unified_ledger(request.user, from_date=f, to_date=t, only_maint_interior=True)
        total = sum([(r.get("credit") or 0) - (r.get("debit") or 0) for r in rows], Decimal("0"))
        return Response({"from": str(f), "to": str(t), "ytd_total": str(total)})


class MIExpensesEntitiesView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        f = parse_date(request.GET.get("from")) or date(date.today().year, 1, 1)
        t = parse_date(request.GET.get("to")) or date.today()
        rows = unified_ledger(request.user, from_date=f, to_date=t, only_maint_interior=True)

        # group by (entity_id, entity_name)
        agg = {}
        for r in rows:
            key = (r.get("entity_id"), r.get("entity") or "—")
            agg[key] = agg.get(key, Decimal("0")) + (r.get("credit") or 0) - (r.get("debit") or 0)

        data = [{"id": k[0], "entity": k[1], "balance": v} for k, v in agg.items()]
        data.sort(key=lambda x: (x["entity"] or ""))
        return Response(EntityBalanceSerializer(data, many=True).data)


class MIExpensesTransactionsView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        f = parse_date(request.GET.get("from")) or date(date.today().year, 1, 1)
        t = parse_date(request.GET.get("to")) or date.today()
        entity_id = request.GET.get("entity_id")
        rows = unified_ledger(
            request.user,
            from_date=f,
            to_date=t,
            entity_id=int(entity_id) if entity_id else None,
            only_maint_interior=True,
        )
        rows = running_balance(rows, opening_balance=Decimal("0"))
        return Response(rows)


class MIExpensesExportView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        f = parse_date(request.GET.get("from")) or date(date.today().year, 1, 1)
        t = parse_date(request.GET.get("to")) or date.today()
        rows = unified_ledger(request.user, from_date=f, to_date=t, only_maint_interior=True)

        agg = {}
        for r in rows:
            key = r.get("entity") or "—"
            agg[key] = agg.get(key, Decimal("0")) + (r.get("credit") or 0) - (r.get("debit") or 0)

        xlsx = export_excel(["Entity", "Balance"], [[k, v] for k, v in agg.items()])
        resp = HttpResponse(
            xlsx.read(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        resp["Content-Disposition"] = f'attachment; filename="mi_entity_balance_{f}_{t}.xlsx"'
        return resp


# --------------------------- Report 3: Owner Dashboard – Rental ---------------------------

def _get_property_model():
    try:
        return apps.get_model("properties", "Property")
    except Exception:
        return None


def _extract_entity_id_from_property(p):
    """
    Try hard to find an entity id for the statement:
    - p.entity_id or p.entity.id
    - p.landlord / p.owner / p.tenant (and their .entity/.id)
    - p.unit / p.apartment (and their .entity/.id)
    """
    # direct FK
    eid = getattr(p, "entity_id", None)
    if eid:
        return eid
    ent = getattr(p, "entity", None)
    if ent:
        return getattr(ent, "id", None) or getattr(ent, "entity_id", None)

    # landlord / owner / tenant
    for who in ("landlord", "owner", "tenant", "tenant_entity", "owner_entity"):
        obj = getattr(p, who, None)
        if obj:
            eid = getattr(obj, "entity_id", None) or getattr(obj, "id", None)
            if eid:
                return eid
            inner = getattr(obj, "entity", None)
            if inner:
                eid = getattr(inner, "id", None) or getattr(inner, "entity_id", None)
                if eid:
                    return eid

    # unit / apartment → entity
    for uni in ("unit", "apartment"):
        obj = getattr(p, uni, None)
        if obj:
            eid = getattr(obj, "entity_id", None)
            if eid:
                return eid
            inner = getattr(obj, "entity", None)
            if inner:
                eid = getattr(inner, "id", None) or getattr(inner, "entity_id", None)
                if eid:
                    return eid
    return None


class OwnerRentalSummaryView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        Property = _get_property_model()
        if Property is None:
            return Response(
                OwnerRentalSummarySerializer(
                    {
                        "total_properties": 0,
                        "rented": 0,
                        "vacant": 0,
                        "care": 0,
                        "sale": 0,
                        "expected_rent_this_month": "0",
                        "igen_sc_this_month": "0",
                        "inspections_30d": 0,
                        "to_be_vacated_30d": 0,
                    }
                ).data
            )

        qs = Property.objects.all()
        # Basic scoping by user's companies if available
        user = request.user
        companies_rel = getattr(user, "companies", None)
        if getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
            try:
                qs = qs.filter(company__in=companies_rel.all())
            except Exception:
                pass

        rented = qs.filter(status__iexact="occupied").count() if hasattr(Property, "status") else 0
        vacant = qs.filter(status__iexact="vacant").count() if hasattr(Property, "status") else 0
        care = qs.filter(purpose__iexact="care").count() if hasattr(Property, "purpose") else 0
        sale = qs.filter(purpose__iexact="sale").count() if hasattr(Property, "purpose") else 0

        # Prefer monthly_rent; fall back to expected_rent
        rent_field = "monthly_rent" if hasattr(Property, "monthly_rent") else ("expected_rent" if hasattr(Property, "expected_rent") else None)
        exp_rent = 0
        if rent_field:
            exp_rent = qs.filter(status__iexact="occupied").aggregate(x=Sum(rent_field)).get("x") or 0

        sc_field = "igen_service_charge" if hasattr(Property, "igen_service_charge") else None
        igen_sc = 0
        if sc_field:
            igen_sc = qs.filter(status__iexact="occupied").aggregate(x=Sum(sc_field)).get("x") or 0

        today = date.today()
        in_30 = today + timedelta(days=30)
        inspections_30d = 0  # (no inspection model; keep 0)
        lease_end_field = "lease_end_date" if hasattr(Property, "lease_end_date") else ("lease_expiry" if hasattr(Property, "lease_expiry") else None)
        to_be_vacated_30d = 0
        if lease_end_field:
            to_be_vacated_30d = qs.filter(**{f"{lease_end_field}__lte": in_30, f"{lease_end_field}__gte": today}).count()

        payload = {
            "total_properties": qs.count(),
            "rented": rented,
            "vacant": vacant,
            "care": care,
            "sale": sale,
            "expected_rent_this_month": str(exp_rent),
            "igen_sc_this_month": str(igen_sc),
            "inspections_30d": inspections_30d,
            "to_be_vacated_30d": to_be_vacated_30d,
        }
        return Response(OwnerRentalSummarySerializer(payload).data)


class OwnerRentalPropertiesView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        Property = _get_property_model()
        if Property is None:
            return Response([])

        qs = Property.objects.all()
        # scope by company if needed
        user = request.user
        companies_rel = getattr(user, "companies", None)
        if getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
            try:
                qs = qs.filter(company__in=companies_rel.all())
            except Exception:
                pass

        rows = []
        for p in qs:
            # Prefer monthly_rent -> expected_rent
            rent_val = getattr(p, "monthly_rent", None)
            if rent_val is None:
                rent_val = getattr(p, "expected_rent", 0)

            # Tenant / owner display via Contact FKs (landlord/tenant)
            tenant_name = getattr(getattr(p, "tenant", None), "full_name", None)
            if not tenant_name:
                tenant_name = getattr(getattr(p, "tenant", None), "name", None)
            owner_name = getattr(getattr(p, "landlord", None), "full_name", None) or getattr(getattr(p, "landlord", None), "name", None)

            rows.append(
                {
                    "id": p.id,
                    "property_name": getattr(p, "name", f"Property {p.id}"),
                    "status": getattr(p, "status", "Vacant"),
                    "rent": rent_val or 0,
                    "igen_service_charge": getattr(p, "igen_service_charge", 0) or 0,
                    "lease_start": getattr(p, "lease_start_date", None) or getattr(p, "lease_start", None),
                    "lease_expiry": getattr(p, "lease_end_date", None) or getattr(p, "lease_expiry", None),
                    "agreement_renewal_date": None,  # not in current model
                    "inspection_date": getattr(p, "next_inspection_date", None) or getattr(p, "inspection_date", None),
                    "tenant_or_owner": tenant_name or owner_name,
                    "transaction_scheduled": getattr(p, "transaction_scheduled", False),
                    "email_sent": getattr(p, "email_sent", False),
                    # ✅ robust entity detection for Generate Statement
                    "entity_id": _extract_entity_id_from_property(p),
                }
            )
        return Response(OwnerRentalRowSerializer(rows, many=True).data)


class OwnerRentalPropertyPatchView(APIView):
    permission_classes = [permissions.IsAuthenticated]


    def patch(self, request, pk):
        from django.shortcuts import get_object_or_404
        from rest_framework.response import Response
        from rest_framework import status

        obj = get_object_or_404(Property, pk=pk)
        data = request.data or {}

        def to_bool(v):
            if isinstance(v, bool):
                return v
            if isinstance(v, str):
                return v.strip().lower() in ("1", "true", "t", "yes", "y", "on")
            try:
                return bool(int(v))
            except Exception:
                return False

        changed = {}

        # Accept both keys for txn scheduled (FE may send either)
        if ("transaction_scheduled" in data) or ("txn_scheduled" in data):
            v = data.get("transaction_scheduled", data.get("txn_scheduled"))
            obj.transaction_scheduled = to_bool(v)
            changed["transaction_scheduled"] = obj.transaction_scheduled

        if "email_sent" in data:
            obj.email_sent = to_bool(data["email_sent"])
            changed["email_sent"] = obj.email_sent

        if not changed:
            return Response({"status": "no-op"}, status=status.HTTP_200_OK)

        obj.save(update_fields=list(changed.keys()))
        return Response({"status": "ok", "applied": changed}, status=status.HTTP_200_OK)

        def to_bool(v):
            if isinstance(v, bool):
                return v
            if isinstance(v, str):
                return v.strip().lower() in ("1","true","t","yes","y","on")
            try:
                return bool(int(v))
            except Exception:
                return False

        updates = {}

        # Accept both keys; alias takes effect if provided
        if "transaction_scheduled" in data:
            updates["transaction_scheduled"] = to_bool(data.get("transaction_scheduled"))
        if "txn_scheduled" in data:
            updates["transaction_scheduled"] = to_bool(data.get("txn_scheduled"))

        if "email_sent" in data:
            updates["email_sent"] = to_bool(data.get("email_sent"))

        # Apply updates if any
        for k, v in updates.items():
            setattr(obj, k, v)
        if updates:
            obj.save()

        return Response({"status": "ok", "applied": updates}, status=status.HTTP_200_OK)


class ProjectProfitabilitySummaryView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        f = parse_date(request.GET.get("from")) or date(date.today().year, 1, 1)
        t = parse_date(request.GET.get("to")) or date.today()
        project_id = request.GET.get("project_id")

        rows = unified_ledger(
            request.user,
            from_date=f,
            to_date=t,
            project_id=int(project_id) if project_id else None,
        )

        agg = {}
        for r in rows:
            proj = r.get("entity") or r.get("contract") or "—"  # adapt as needed
            a = agg.setdefault(proj, {"in": Decimal("0"), "out": Decimal("0")})
            a["in"] += r.get("credit") or 0
            a["out"] += r.get("debit") or 0

        out = [
            {"project": k, "inflows": v["in"], "outflows": v["out"], "net": v["in"] - v["out"]}
            for k, v in agg.items()
        ]
        out.sort(key=lambda x: x["project"])
        return Response(ProjectProfitRowSerializer(out, many=True).data)


class ProjectProfitabilityTransactionsView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        f = parse_date(request.GET.get("from")) or date(date.today().year, 1, 1)
        t = parse_date(request.GET.get("to")) or date.today()
        project_id = request.GET.get("project_id")

        rows = unified_ledger(
            request.user,
            from_date=f,
            to_date=t,
            project_id=int(project_id) if project_id else None,
        )
        rows = running_balance(rows, opening_balance=Decimal("0"))
        return Response(rows)


class ProjectProfitabilityExportView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        f = parse_date(request.GET.get("from")) or date(date.today().year, 1, 1)
        t = parse_date(request.GET.get("to")) or date.today()
        rows = ProjectProfitabilitySummaryView().get(request).data  # reuse payload

        xlsx = export_excel(
            ["Project", "Inflows", "Outflows", "Net"],
            [[r["project"], r["inflows"], r["outflows"], r["net"]] for r in rows],
        )
        resp = HttpResponse(
            xlsx.read(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        resp["Content-Disposition"] = f'attachment; filename="project_profitability_{f}_{t}.xlsx"'
        return resp


# --------------------------- Report 5: Financial Dashboard (Pivot) ---------------------------

class FinancialDashboardPivotView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        body = request.data or {}
        dims = body.get("dims") or []  # e.g. ["cost_centre","txn_type","entity","asset","contract","date"]
        values = body.get("values") or {}  # {"entity":[...], ...}
        date_gran = body.get("date_granularity") or body.get("date_gran") or body.get("granularity") or "day"
        f = parse_date(body.get("from")) or date(date.today().year, 1, 1)
        t = parse_date(body.get("to")) or date.today()

        rows = unified_ledger(request.user, from_date=f, to_date=t)

        # filter
        def allowed(r):
            for d in dims:
                sel = values.get(d)
                if d == "date":
                    val = _format_period(r.get("value_date"), date_gran)
                else:
                    val = _dim_value(r, d) or "—"
                if sel and val not in set(sel):
                    return False
            return True

        filt = [r for r in rows if allowed(r)]

        # group by dims
        from collections import defaultdict
        grp = defaultdict(lambda: {"credit": Decimal("0"), "debit": Decimal("0")})
        for r in filt:
            key_items = []
            for d in dims:
                if d == "date":
                    key_items.append(_format_period(r.get("value_date"), date_gran))
                else:
                    key_items.append((_dim_value(r, d) or "—"))
            key = tuple(key_items)
            grp[key]["credit"] += r.get("credit") or 0
            grp[key]["debit"] += r.get("debit") or 0

        out = []
        for key, v in grp.items():
            item = {dims[i]: key[i] for i in range(len(dims))}
            item["credit"] = v["credit"]
            item["debit"] = v["debit"]
            item["margin"] = v["credit"] - v["debit"]
            out.append(item)

        # running balance across filtered set (by value date)
        bal = Decimal("0")
        for r in sorted(filt, key=lambda x: x["value_date"]):
            bal += (r.get("credit") or 0) - (r.get("debit") or 0)

        totals = {
            "credit": sum([i["credit"] for i in out], Decimal("0")),
            "debit": sum([i["debit"] for i in out], Decimal("0")),
            "margin": sum([i["margin"] for i in out], Decimal("0")),
            "balance": bal,
        }
        return Response({"rows": out, "totals": totals})


class FinancialDashboardExportView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        payload = FinancialDashboardPivotView().post(request).data
        rows = payload.get("rows", [])
        headers = sorted({k for r in rows for k in r.keys()})
        xlsx = export_excel(headers, [[r.get(h, "") for h in headers] for r in rows])
        resp = HttpResponse(
            xlsx.read(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        resp["Content-Disposition"] = 'attachment; filename="financial_dashboard_pivot.xlsx"'
        return resp
