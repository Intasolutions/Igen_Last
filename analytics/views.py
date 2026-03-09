from properties.models import Property
from datetime import date, timedelta, datetime
from decimal import Decimal
from io import BytesIO  # for DOCX response
from collections import defaultdict  # needed by FinancialDashboardPivotView

from django.apps import apps
from django.db.models import Sum, Q
from django.http import HttpResponse
from django.utils.dateparse import parse_date
from rest_framework import permissions
from rest_framework.response import Response
from rest_framework.views import APIView
from tx_classify.models import Classification
from analytics.models import OwnerRentalFlag

# Optional Word export (python-docx)
try:
    from docx import Document  # pip install python-docx
    from docx.shared import Pt

    _DOCX_OK = True
except Exception:
    Document = None
    Pt = None
    _DOCX_OK = False

from .serializers import (
    EntityStatementRowSerializer,
    EntityBalanceSerializer,
    OwnerRentalSummarySerializer,
    OwnerRentalRowSerializer,
    ProjectProfitRowSerializer,
    OwnerRentalPendingPropertySerializer,
    OwnerRentalInspectionExpiryPropertySerializer,
    OwnerRentalAgreementExpiryPropertySerializer,
    OwnerRentalServiceChargeBreakdownSerializer,
    OwnerRentalMaintenanceBreakdownSerializer,
)
from .services.exports import export_excel, export_simple_pdf
from .services.ledger import unified_ledger, running_balance, opening_balance_until


# --------------------------- helpers ---------------------------


def _month_range(yyyy_mm: str):
    """
    Return (start, end) where end is the **inclusive** last day of the month.
    """
    y, m = yyyy_mm.split("-")
    y = int(y)
    m = int(m)
    start = date(y, m, 1)
    next_first = date(y + (m == 12), 1 if m == 12 else m + 1, 1)
    end = next_first - timedelta(days=1)
    return start, end

def _months_between(from_yyyy_mm: str, to_yyyy_mm: str):
    """
    Returns a list of YYYY-MM strings from from_yyyy_mm to to_yyyy_mm inclusive.
    Example: 2026-01 to 2026-03 => ["2026-01","2026-02","2026-03"]
    """
    fy, fm = map(int, from_yyyy_mm.split("-"))
    ty, tm = map(int, to_yyyy_mm.split("-"))

    start = date(fy, fm, 1)
    end = date(ty, tm, 1)

    # if user selects reverse order, swap
    if start > end:
        start, end = end, start

    out = []
    cur = start
    while cur <= end:
        out.append(f"{cur.year:04d}-{cur.month:02d}")
        # increment month safely (no extra libs)
        if cur.month == 12:
            cur = date(cur.year + 1, 1, 1)
        else:
            cur = date(cur.year, cur.month + 1, 1)
    return out


def _statement_5th_range(yyyy_mm: str):
    """
    Returns (from_date, to_date) where:
    - from_date = previous month 5th
    - to_date   = current month 5th
    Example:
      2025-12 -> 2025-11-05 to 2025-12-05
      2026-01 -> 2025-12-05 to 2026-01-05
    """
    y, m = map(int, yyyy_mm.split("-"))

    # current month 5th
    to_date = date(y, m, 5)

    # previous month 5th
    if m == 1:
        from_date = date(y - 1, 12, 5)
    else:
        from_date = date(y, m - 1, 5)

    return from_date, to_date


def _dim_value(row, dim: str):
    """
    Read a dimension out of a unified_ledger row.
    Special-case 'date' to use the value_date.
    """
    if dim == "date":
        return row.get("value_date")
    return row.get(dim)


def _format_period(d: date, granularity: str) -> str:
    """
    Map a date to day/month/quarter/year label strings for pivot 'date' dim.
    """
    if not d:
        return "—"
    g = (granularity or "day").lower()
    if g == "day":
        return d.isoformat()
    if g == "month":
        return f"{d.year:04d}-{d.month:02d}"
    if g == "quarter":
        q = (d.month - 1) // 3 + 1
        return f"{d.year:04d}-Q{q}"
    if g == "year":
        return f"{d.year:04d}"
    return d.isoformat()


def _safe_filename_part(s) -> str:
    """
    Make a string safe to use in filenames (no spaces/special chars).
    """
    txt = str(s or "").strip()
    if not txt:
        return "NA"
    return "".join(ch if ch.isalnum() or ch in ("-", "_") else "_" for ch in txt)


def _get_entity_display(entity_id: int):
    """
    Helper to get a nice label for an Entity:
      - name
      - optional code/property_code
    """
    try:
        Entity = apps.get_model("entities", "Entity")
    except Exception:
        return {"name": f"Entity {entity_id}", "code": None}

    try:
        ent = Entity.objects.get(pk=entity_id)
    except Exception:
        return {"name": f"Entity {entity_id}", "code": None}

    name = getattr(ent, "name", None) or f"Entity {entity_id}"
    code = (
        getattr(ent, "code", None)
        or getattr(ent, "property_code", None)
        or getattr(ent, "short_code", None)
    )
    return {"name": name, "code": code}


# --------------------------- Quick health/debug ---------------------------


class AnalyticsHealthView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        # very wide range just to see *something*
        rows = unified_ledger(
            request.user,
            from_date=date(2000, 1, 1),
            to_date=date(2100, 1, 1),
        )
        return Response({"ok": True, "row_count": len(rows)})


class AnalyticsDataProbeView(APIView):
    """
    Debug endpoint: shows discovered models & counts for cash_ledger, tx_classify, bank_uploads.
    """

    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        def model_info(app_label):
            out = []
            try:
                cfg = apps.get_app_config(app_label)
            except Exception:
                return {"error": "app not installed"}
            for M in cfg.get_models():
                try:
                    fields = [f.name for f in M._meta.get_fields()]
                    # pick a likely date field
                    date_field = next(
                        (
                            d
                            for d in [
                                "value_date",
                                "transaction_date",
                                "date",
                                "posting_date",
                                "book_date",
                            ]
                            if d in fields
                        ),
                        None,
                    )
                    cnt = M.objects.count()
                    sample = []
                    if cnt:
                        sample = list(M.objects.values_list(date_field or "id")[:3])
                    out.append(
                        {
                            "model": M.__name__,
                            "count": cnt,
                            "date_field": date_field,
                            "has_credit": "credit" in fields,
                            "has_debit": "debit" in fields,
                            "has_amount": any(
                                a in fields
                                for a in ["amount", "value", "deposit", "withdrawal"]
                            ),
                            "fields": fields[:30],
                            "sample": sample,
                        }
                    )
                except Exception as e:
                    out.append({"model": M.__name__, "error": str(e)})
            return out

        payload = {
            "cash_ledger": model_info("cash_ledger"),
            "tx_classify": model_info("tx_classify"),
            "bank_uploads": model_info("bank_uploads"),
        }
        return Response(payload)


# --------------------------- NEW: Entity quick search ---------------------------


class EntityQuickSearchView(APIView):
    """
    Read-only helper for the UI to discover entity IDs safely.
    GET params:
      - q: search text (matches name icontains or numeric id)
      - limit: optional (default 20, max 50)
    Response: [{id, name}]
    """

    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        q = (request.GET.get("q") or "").strip()
        try:
            limit = min(max(int(request.GET.get("limit", "20")), 1), 50)
        except Exception:
            limit = 20

        try:
            Entity = apps.get_model("entities", "Entity")
        except Exception:
            return Response([], status=200)

        qs = Entity.objects.all().order_by("name")
        if q:
            cond = Q(name__icontains=q)
            if q.isdigit():
                cond |= Q(id=int(q))
            qs = qs.filter(cond)

        rows = list(qs.values("id", "name")[:limit])
        return Response(rows)


# --------------------------- Report 1: Entity-wise Monthly Statement ---------------------------


class EntityStatementView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        entity_id = request.GET.get("entity_id")
        month = request.GET.get("month")        # old: YYYY-MM
        from_str = request.GET.get("from")      # new: YYYY-MM-DD
        to_str = request.GET.get("to")          # new: YYYY-MM-DD

        if not entity_id:
            return Response({"detail": "entity_id required"}, status=400)

        # ---------- MODE 1: old single-month ----------
        if month:
            start, end = _month_range(month)

        # ---------- MODE 2: new date-range ----------
        elif from_str and to_str:
            from_date = parse_date(from_str)
            to_date = parse_date(to_str)

            if not (from_date and to_date):
                return Response(
                    {"detail": "from & to must be valid dates (YYYY-MM-DD)."},
                    status=400,
                )

            # swap if reversed
            if from_date > to_date:
                from_date, to_date = to_date, from_date

            start, end = from_date, to_date

        else:
            return Response(
                {"detail": "Provide either month (YYYY-MM) OR from & to (YYYY-MM-DD)."},
                status=400,
            )

        # Opening balance strictly before start
        obal = opening_balance_until(request.user, start, entity_id=int(entity_id))

        base_rows = unified_ledger(
            request.user,
            from_date=start,
            to_date=end,
            entity_id=int(entity_id),
        )

        rows = running_balance(base_rows, opening_balance=obal)

        data = [
            {
                "value_date": r["value_date"],
                "txn_type": r.get("txn_type"),
                "credit": r.get("credit", Decimal("0")),
                "debit": r.get("debit", Decimal("0")),
                "balance": r.get("balance", Decimal("0")),
                "remarks": r.get("remarks") or "",
            }
            for r in rows
        ]
        return Response(EntityStatementRowSerializer(data, many=True).data)




class EntityStatementPDFView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        entity_id = request.GET.get("entity_id")
        month = request.GET.get("month")        # old: YYYY-MM
        from_str = request.GET.get("from")      # new: YYYY-MM-DD
        to_str = request.GET.get("to")          # new: YYYY-MM-DD

        if not entity_id:
            return Response({"detail": "entity_id required"}, status=400)

        # Entity display
        try:
            eid_int = int(entity_id)
        except Exception:
            eid_int = None

        if eid_int is not None:
            ent_label = _get_entity_display(eid_int)
            ent_name = ent_label["name"]
            ent_code = ent_label["code"]
        else:
            ent_name = f"Entity {entity_id}"
            ent_code = None

        # ----- period selection -----
        if month:
            start, end = _month_range(month)
            period_label = month

        elif from_str and to_str:
            from_date = parse_date(from_str)
            to_date = parse_date(to_str)

            if not (from_date and to_date):
                return Response(
                    {"detail": "from & to must be valid dates (YYYY-MM-DD)."},
                    status=400,
                )

            if from_date > to_date:
                from_date, to_date = to_date, from_date

            start, end = from_date, to_date
            period_label = f"{start}_to_{end}"

        else:
            return Response(
                {"detail": "Provide either month (YYYY-MM) OR from & to (YYYY-MM-DD)."},
                status=400,
            )

        # Opening balance before start
        obal = opening_balance_until(request.user, start, entity_id=int(entity_id))

        base_rows = unified_ledger(
            request.user,
            from_date=start,
            to_date=end,
            entity_id=int(entity_id),
        )
        rows = running_balance(base_rows, opening_balance=obal)

        headers = ["Date", "Type", "Credit", "Debit", "Balance", "Remarks"]
        table = [
            [
                r["value_date"].strftime("%Y-%m-%d") if r.get("value_date") else "",
                r.get("txn_type") or "",
                str(r.get("credit") or 0),
                str(r.get("debit") or 0),
                str(r.get("balance") or 0),
                (r.get("remarks") or "")[:64],
            ]
            for r in rows
        ]

        # Title
        if ent_code:
            title = f"{ent_name} - {ent_code} - {period_label} Statement"
        else:
            title = f"{ent_name} - {period_label} Statement"

        pdf = export_simple_pdf(title, headers, table)
        resp = HttpResponse(pdf.read(), content_type="application/pdf")
        resp["Content-Disposition"] = (
            f'attachment; filename="entity_statement_{_safe_filename_part(ent_name)}_{period_label}.pdf"'
        )
        return resp




class EntityStatementDOCXView(APIView):
    """
    Word export for Entity Statement.
    Supports:
      - entity_id + month (YYYY-MM)   [old]
      - entity_id + from/to (YYYY-MM-DD) [new]
    """

    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        if not _DOCX_OK:
            return Response(
                {"detail": "Word export unavailable: install python-docx"},
                status=500,
            )

        entity_id = request.GET.get("entity_id")
        month = request.GET.get("month")        # old
        from_str = request.GET.get("from")      # new
        to_str = request.GET.get("to")          # new

        if not entity_id:
            return Response({"detail": "entity_id required"}, status=400)

        # entity label
        try:
            eid_int = int(entity_id)
        except Exception:
            eid_int = None

        if eid_int is not None:
            ent_label = _get_entity_display(eid_int)
            ent_name = ent_label["name"]
            ent_code = ent_label.get("code")
        else:
            ent_name = f"Entity {entity_id}"
            ent_code = None

        # ----- period selection -----
        if month:
            start, end = _month_range(month)
            period_label = month

        elif from_str and to_str:
            from_date = parse_date(from_str)
            to_date = parse_date(to_str)

            if not (from_date and to_date):
                return Response(
                    {"detail": "from & to must be valid dates (YYYY-MM-DD)."},
                    status=400,
                )

            if from_date > to_date:
                from_date, to_date = to_date, from_date

            start, end = from_date, to_date
            period_label = f"{start}_to_{end}"

        else:
            return Response(
                {"detail": "Provide either month (YYYY-MM) OR from & to (YYYY-MM-DD)."},
                status=400,
            )

        # data
        obal = opening_balance_until(request.user, start, entity_id=int(entity_id))
        base_rows = unified_ledger(
            request.user,
            from_date=start,
            to_date=end,
            entity_id=int(entity_id),
        )
        rows = running_balance(base_rows, opening_balance=obal)

        # build docx
        doc = Document()
        title = f"{ent_name} - Statement"
        if ent_code:
            title = f"{ent_name} - {ent_code} - Statement"
        doc.add_heading(title, level=1)

        p = doc.add_paragraph()
        p.add_run(f"Period: {period_label}").italic = True

        headers = ["Date", "Type", "Credit", "Debit", "Balance", "Remarks"]
        table = doc.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            run = hdr_cells[i].paragraphs[0].add_run(h)
            if Pt:
                run.font.bold = True
                run.font.size = Pt(10)

        for r in rows:
            row_cells = table.add_row().cells
            row_cells[0].text = (
                r["value_date"].strftime("%Y-%m-%d") if r.get("value_date") else ""
            )
            row_cells[1].text = str(r.get("txn_type") or "")
            row_cells[2].text = str(r.get("credit") or 0)
            row_cells[3].text = str(r.get("debit") or 0)
            row_cells[4].text = str(r.get("balance") or 0)
            row_cells[5].text = (r.get("remarks") or "")[:256]

        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)

        resp = HttpResponse(
            bio.getvalue(),
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )
        resp["Content-Disposition"] = (
            f'attachment; filename="entity_statement_{_safe_filename_part(ent_name)}_{period_label}.docx"'
        )
        return resp



class EntityStatementExcelView(APIView):
    """
    Excel export for Entity Statement (Report 1).
    """

    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        entity_id = request.GET.get("entity_id")
        month = request.GET.get("month")
        if not (entity_id and month):
            return Response(
                {"detail": "entity_id & month (YYYY-MM) required"},
                status=400,
            )

        try:
            eid_int = int(entity_id)
        except Exception:
            eid_int = None

        if eid_int is not None:
            ent_label = _get_entity_display(eid_int)
            ent_name = ent_label["name"]
        else:
            ent_name = f"Entity {entity_id}"

        start, end = _month_range(month)
        obal = opening_balance_until(request.user, start, entity_id=int(entity_id))
        base_rows = unified_ledger(
            request.user,
            from_date=start,
            to_date=end,
            entity_id=int(entity_id),
        )
        rows = running_balance(base_rows, opening_balance=obal)

        headers = ["Date", "Type", "Credit", "Debit", "Balance", "Remarks"]
        data = [
            [
                r["value_date"].strftime("%Y-%m-%d"),
                r.get("txn_type") or "",
                r.get("credit") or 0,
                r.get("debit") or 0,
                r.get("balance") or 0,
                (r.get("remarks") or "")[:256],
            ]
            for r in rows
        ]

        xlsx = export_excel(headers, data)
        resp = HttpResponse(
            xlsx.read(),
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "spreadsheetml.sheet"
            ),
        )
        resp[
            "Content-Disposition"
        ] = f'attachment; filename="entity_statement_{_safe_filename_part(ent_name)}_{month}.xlsx"'
        return resp


# --------------------------- Report 2: Maintenance & Interior (YTD) ---------------------------


class MIExpensesSummaryView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        f = parse_date(request.GET.get("from")) or date(date.today().year, 1, 1)
        t = parse_date(request.GET.get("to")) or date.today()
        rows = unified_ledger(
            request.user,
            from_date=f,
            to_date=t,
            only_maint_interior=True,
        )
        # For spend, treat DEBIT as positive, CREDIT as offset
        total = sum(
            [(r.get("debit") or 0) - (r.get("credit") or 0) for r in rows],
            Decimal("0"),
        )
        return Response({"from": str(f), "to": str(t), "ytd_total": str(total)})


class MIExpensesEntitiesView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        f = parse_date(request.GET.get("from")) or date(date.today().year, 1, 1)
        t = parse_date(request.GET.get("to")) or date.today()
        rows = unified_ledger(
            request.user,
            from_date=f,
            to_date=t,
            only_maint_interior=True,
        )

        # group by (entity_id, entity_name)
        agg = {}
        for r in rows:
            key = (r.get("entity_id"), r.get("entity") or "—")
            # debit (spend) minus credit (reversals/adjustments)
            agg[key] = agg.get(key, Decimal("0")) + (r.get("debit") or 0) - (
                r.get("credit") or 0
            )

        data = [{"id": k[0], "entity": k[1], "balance": v} for k, v in agg.items()]
        data.sort(key=lambda x: (x["entity"] or ""))
        return Response(EntityBalanceSerializer(data, many=True).data)


class MIExpensesTransactionsView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        f = parse_date(request.GET.get("from")) or date(date.today().year, 1, 1)
        t = parse_date(request.GET.get("to")) or date.today()
        entity_id = request.GET.get("entity_id")
        rows = unified_ledger(
            request.user,
            from_date=f,
            to_date=t,
            entity_id=int(entity_id) if entity_id else None,
            only_maint_interior=True,
        )
        rows = running_balance(rows, opening_balance=Decimal("0"))
        return Response(rows)


class MIExpensesExportView(APIView):
    """
    Excel export that matches the entity-wise M&I YTD view:
      - one row per (entity_id, entity_name)
      - balance = debit - credit
    """

    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        f = parse_date(request.GET.get("from")) or date(date.today().year, 1, 1)
        t = parse_date(request.GET.get("to")) or date.today()
        rows = unified_ledger(
            request.user,
            from_date=f,
            to_date=t,
            only_maint_interior=True,
        )

        agg = {}
        for r in rows:
            key = (r.get("entity_id"), r.get("entity") or "—")
            agg[key] = agg.get(key, Decimal("0")) + (r.get("debit") or 0) - (
                r.get("credit") or 0
            )

        headers = ["Entity ID", "Entity", "Balance"]
        data = [[k[0], k[1], v] for k, v in agg.items()]

        xlsx = export_excel(headers, data)
        resp = HttpResponse(
            xlsx.read(),
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "spreadsheetml.sheet"
            ),
        )
        resp[
            "Content-Disposition"
        ] = f'attachment; filename="mi_ytd_entity_balance_{f}_{t}.xlsx"'
        return resp


# --------------------------- Report 3: Owner Dashboard – Rental ---------------------------


def _get_property_model():
    try:
        return apps.get_model("properties", "Property")
    except Exception:
        return None


def _entity_for_property(prop: Property):
    """
    Resolve the Entity row linked to a Property via entities.Entity.linked_property.

    We keep this SAFE for production:
      - Always filter by linked_property.
      - If fields like entity_type/status exist, apply nice filters.
      - Otherwise just pick the latest row.
    """
    try:
        Entity = apps.get_model("entities", "Entity")
    except Exception:
        return None

    qs = Entity.objects.filter(linked_property=prop)

    # Optional filters – only if those fields actually exist on the model.
    field_names = {f.name for f in Entity._meta.get_fields()}

    if "entity_type" in field_names:
        qs = qs.filter(Q(entity_type="Property") | Q(entity_type__iexact="property"))

    if "status" in field_names:
        qs = qs.filter(Q(status="Active") | Q(status__iexact="active"))

    # Prefer created_at if present, else id
    if "created_at" in field_names:
        qs = qs.order_by("-created_at")
    else:
        qs = qs.order_by("-id")

    try:
        return qs.first()
    except Exception:
        return None


def _extract_entity_id_from_property(p):
    """
    Try to find an entity id for the statement:
    - Entity linked via entities.Entity.linked_property (preferred)
    - p.entity_id or p.entity.id
    - contacts linked via tenant_contact / landlord / owner / tenant_entity / owner_entity
      (and their .entity/.entity_id)
    - p.unit / p.apartment (and their .entity/.id)
    """
    # First, prefer the canonical Entity linkage
    ent = _entity_for_property(p)
    if ent:
        return getattr(ent, "id", None) or getattr(ent, "entity_id", None)

    # direct FK if ever added
    eid = getattr(p, "entity_id", None)
    if eid:
        return eid
    ent = getattr(p, "entity", None)
    if ent:
        return getattr(ent, "id", None) or getattr(ent, "entity_id", None)

    # landlord / tenant contacts
    for who in ("tenant_contact", "landlord", "owner", "tenant_entity", "owner_entity"):
        obj = getattr(p, who, None)
        if not obj:
            continue
        inner = getattr(obj, "entity", None)
        if inner:
            eid = getattr(inner, "id", None) or getattr(inner, "entity_id", None)
            if eid:
                return eid
        # if the contact itself carries an entity_id field
        eid = getattr(obj, "entity_id", None)
        if eid:
            return eid

    # unit / apartment → entity
    for uni in ("unit", "apartment"):
        obj = getattr(p, uni, None)
        if obj:
            eid = getattr(obj, "entity_id", None)
            if eid:
                return eid
            inner = getattr(obj, "entity", None)
            if inner:
                eid = getattr(inner, "id", None) or getattr(inner, "entity_id", None)
                if eid:
                    return eid
    return None


def _statement_rows_for_property(user, prop: Property, from_date: date, to_date: date):
    """
    Build a classified, ledger-backed statement for a single property
    over an arbitrary date range [from_date, to_date].

    Implementation:
      - Resolve entities.Entity for this property via linked_property.
      - Pull rows from unified_ledger() filtered by that entity_id and date range.
      - Apply running_balance() with opening_balance_until() as of from_date.
    """
    if not (from_date and to_date):
        return []

    ent = _entity_for_property(prop)
    if not ent:
        return []

    entity_id = getattr(ent, "id", None) or getattr(ent, "entity_id", None)
    if not entity_id:
        return []

    # Opening balance is strictly before from_date
    obal = opening_balance_until(user, from_date, entity_id=entity_id)

    base_rows = unified_ledger(
        user,
        from_date=from_date,
        to_date=to_date,
        entity_id=entity_id,
    )
    rows = running_balance(base_rows, opening_balance=obal)
    return rows


def _synthetic_property_statement_rows(prop: Property, month: str):
    """
    Fallback when unified_ledger returns no rows for a property.

    Builds a simple 'contract-based' statement from the Property master:
      - Rent for the month (CREDIT)
      - iGen Service Charge for the month (DEBIT)

    This ensures the owner always gets something from the Owner Dashboard,
    even before a full ledger integration is in place.
    """
    start, _ = _statement_5th_range(month)


    # Pick rent: monthly_rent → expected_rent → rent → 0
    rent = getattr(prop, "monthly_rent", None)
    if rent is None:
        rent = getattr(prop, "expected_rent", None)
    if rent is None:
        rent = getattr(prop, "rent", None)
    rent = rent or Decimal("0")

    # iGen service charge if present
    sc = getattr(prop, "igen_service_charge", None) or Decimal("0")

    rows = []
    balance = Decimal("0")

    # Rent as CREDIT
    if rent:
        balance += rent
        rows.append(
            {
                "value_date": start,
                "txn_type": "Rent",
                "credit": rent,
                "debit": Decimal("0"),
                "balance": balance,
                "remarks": f"Scheduled rent for {month}",
            }
        )

    # Service Charge as DEBIT
    if sc:
        balance -= sc
        rows.append(
            {
                "value_date": start,
                "txn_type": "Service Charge",
                "credit": Decimal("0"),
                "debit": sc,
                "balance": balance,
                "remarks": f"iGen service charge for {month}",
            }
        )

    if not rows:
        # Absolutely nothing configured – still show a friendly note.
        rows.append(
            {
                "value_date": start,
                "txn_type": "",
                "credit": Decimal("0"),
                "debit": Decimal("0"),
                "balance": Decimal("0"),
                "remarks": "No transactions recorded or configured for this period.",
            }
        )

    return rows


class OwnerRentalSummaryView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        PropertyModel = _get_property_model()
        if PropertyModel is None:
            return Response(
                OwnerRentalSummarySerializer(
                    {
                        "total_properties": 0,
                        "rented": 0,
                        "vacant": 0,
                        "care": 0,
                        "sale": 0,
                        "rent_to_be_collected": "0",
                        "igen_sc_this_month": "0",
                        "inspections_30d": 0,
                        "to_be_vacated_30d": 0,
                    }
                ).data
            )

        qs = PropertyModel.objects.filter(is_active=True)
        user = request.user
        companies_rel = getattr(user, "companies", None)
        if getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
            try:
                qs = qs.filter(company__in=companies_rel.all())
            except Exception:
                pass

        # Parse requested month
        month_str = request.GET.get("month") or date.today().strftime("%Y-%m")
        m_start, m_end = _month_range(month_str)
        days_in_month = Decimal(str((m_end - m_start).days + 1))

        # Stats to compute
        rented_contribution_count = 0
        rent_sum = Decimal("0")
        sc_sum = Decimal("0")

        for p in qs:
            # TC-06/08: Strict Status Validation (Strip + Lowercase)
            status_str = (getattr(p, "status", "") or "").strip().lower()
            if status_str != "occupied":
                continue

            # TC-12/13: Monthly Rent check
            base_rent = getattr(p, "monthly_rent", Decimal("0")) or Decimal("0")
            base_sc = getattr(p, "igen_service_charge", Decimal("0")) or Decimal("0")
            
            l_start = getattr(p, "lease_start_date", None)
            l_end = getattr(p, "lease_end_date", None)
            
            # TC-14: Invalid date guard
            if l_start and l_end and l_end < l_start:
                continue

            # Overlap Logic (TC-01 to TC-05, TC-09 to TC-11)
            actual_start = max(l_start, m_start) if l_start else m_start
            actual_end = min(l_end or m_end, m_end)
            
            if actual_end >= actual_start:
                rented_contribution_count += 1
                occupied_days = (actual_end - actual_start).days + 1
                # Formula: (Monthly Rent / Days in Month) * Occupied Days
                rent_sum += (base_rent / days_in_month) * Decimal(str(occupied_days))
                sc_sum += (base_sc / days_in_month) * Decimal(str(occupied_days))

        # Categories for other tiles (Static counts)
        care = qs.filter(purpose__iexact="care").count()
        sale = qs.filter(purpose__iexact="sale").count()
        total = qs.count()

        today = date.today()
        in_30 = today + timedelta(days=30)
        in_5 = today + timedelta(days=5)

        # Lookaheads
        inspections = 0
        if hasattr(PropertyModel, "next_inspection_date"):
            inspections = qs.filter(next_inspection_date__gte=today, next_inspection_date__lte=in_30).count()
        
        inspections_due = 0
        inspections_expired = 0
        if hasattr(PropertyModel, "next_inspection_date"):
            inspections_due = qs.filter(next_inspection_date__gte=today, next_inspection_date__lte=in_5).count()
            inspections_expired = qs.filter(next_inspection_date__lt=today).count()

        to_vacate = 0
        renewals_30d = 0
        agreements_expired = 0
        if hasattr(PropertyModel, "lease_end_date"):
            to_vacate = qs.filter(lease_end_date__gte=today, lease_end_date__lte=in_30).count()
            # Only for Occupied properties as per 11th enhancement
            renewals_30d = qs.filter(status__iexact="occupied", lease_end_date__gte=today, lease_end_date__lte=in_30).count()
            agreements_expired = qs.filter(status__iexact="occupied", lease_end_date__lt=today).count()

        # Rent Received Calculation
        rent_received_filters = Q(
            is_active_classification=True,
            transaction_type__name__iexact='Rent In', # Requirement: Only take 'Rent In'
            value_date__range=(m_start, m_end),
            bank_transaction__source='BANK',
            bank_transaction__is_deleted=False,
            bank_transaction__credit_amount__gt=0
        )
        
        # TC-09: Filter by company access
        if getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
            rent_received_filters &= Q(transaction_type__company__in=companies_rel.all())

        from cash_ledger.models import CashLedgerRegister

        # Owner Recoverables (Maintenance/Expenses)
        # 1. BANK Recoverables
        bank_recoverables_qs = Classification.objects.filter(
            is_active_classification=True,
            transaction_type__name__iregex=r'(Maintenance|Legal|Paper Work|Paperwork)',
            value_date__range=(m_start, m_end),
            bank_transaction__is_deleted=False
        )
        if getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
            bank_recoverables_qs = bank_recoverables_qs.filter(transaction_type__company__in=companies_rel.all())

        # We must iterate or use annotation because margin is in remarks text
        bank_recoverables_base = Decimal('0')
        bank_recoverables_margin = Decimal('0')
        for item in bank_recoverables_qs:
            bank_recoverables_base += (item.amount or Decimal('0'))
            bank_recoverables_margin += (item.parsed_margin or Decimal('0'))
        
        bank_recoverables_sum = bank_recoverables_base + bank_recoverables_margin

        # 2. CASH Recoverables
        cash_recoverables_filters = Q(
            is_active=True,
            transaction_type__name__iregex=r'(Maintenance|Legal|Paper Work|Paperwork)',
            date__range=(m_start, m_end)
        )
        if getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
            cash_recoverables_filters &= Q(company__in=companies_rel.all())

        # In CashLedger, both amount and margin are explicit fields
        cash_agg = CashLedgerRegister.objects.filter(cash_recoverables_filters).aggregate(
            base=Sum('amount'),
            mgn=Sum('margin')
        )
        cash_recoverables_base = (cash_agg['base'] or Decimal('0'))
        cash_recoverables_margin = (cash_agg['mgn'] or Decimal('0'))
        cash_recoverables_sum = cash_recoverables_base + cash_recoverables_margin

        owner_recoverables_sum = bank_recoverables_sum + cash_recoverables_sum
        owner_recoverables_base = bank_recoverables_base + cash_recoverables_base
        owner_recoverables_margin = bank_recoverables_margin + cash_recoverables_margin

        # Total Margin Collected (Requirement 7 Update: Only Margin Applicable Types + Breakdown)
        margin_breakdown = defaultdict(lambda: {"bank": Decimal('0'), "cash": Decimal('0'), "total": Decimal('0')})
        
        # 1. BANK Margin
        bank_margin_qs = Classification.objects.filter(
            is_active_classification=True,
            transaction_type__margin_applicable=True, # New Strict Requirement
            value_date__range=(m_start, m_end),
            bank_transaction__is_deleted=False
        ).select_related('cost_centre')
        
        if getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
            bank_margin_qs = bank_margin_qs.filter(transaction_type__company__in=companies_rel.all())

        bank_margin_sum = Decimal('0')
        for item in bank_margin_qs:
            mgn = item.parsed_margin
            if mgn and mgn > 0:
                cc_name = item.cost_centre.name if item.cost_centre else "Other"
                margin_breakdown[cc_name]["bank"] += mgn
                margin_breakdown[cc_name]["total"] += mgn
                bank_margin_sum += mgn

        # 2. CASH Margin
        cash_margin_filters = Q(
            is_active=True,
            transaction_type__margin_applicable=True, # New Strict Requirement
            date__range=(m_start, m_end),
            margin__gt=0
        )
        if getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
            cash_margin_filters &= Q(company__in=companies_rel.all())

        cash_margin_qs = CashLedgerRegister.objects.filter(cash_margin_filters).select_related('cost_centre')
        cash_margin_sum = Decimal('0')
        
        for item in cash_margin_qs:
            mgn = item.margin or Decimal('0')
            cc_name = item.cost_centre.name if item.cost_centre else "Other"
            margin_breakdown[cc_name]["cash"] += mgn
            margin_breakdown[cc_name]["total"] += mgn
            cash_margin_sum += mgn

        total_margin_collected_sum = bank_margin_sum + cash_margin_sum
        
        # Format breakdown for JSON
        formatted_margin_breakdown = [
            {"cost_centre": cc, **vals} for cc, vals in margin_breakdown.items()
        ]

        rent_received_sum = Classification.objects.filter(rent_received_filters).aggregate(
            total=Sum('amount')
        )['total'] or Decimal('0')

        # iGen Service Charge Collected
        sc_collected_filters = Q(
            is_active_classification=True,
            transaction_type__name__icontains='Service Charge', # Flex search for customizable names
            value_date__range=(m_start, m_end),
            bank_transaction__source='BANK',
            bank_transaction__is_deleted=False,
            bank_transaction__credit_amount__gt=0
        )
        if getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
            sc_collected_filters &= Q(transaction_type__company__in=companies_rel.all())

        sc_collected_sum = Classification.objects.filter(sc_collected_filters).aggregate(
            total=Sum('amount')
        )['total'] or Decimal('0')

        # Total iGen Income (8th Enhancement)
        income_types = ['iGen Service Charge', 'iGen Brokerage', 'Other Income']
        igen_income_type_breakdown = {t: Decimal('0') for t in income_types}
        igen_income_cc_breakdown = defaultdict(Decimal)
        
        # 1. BANK Income Credits
        bank_income_qs = Classification.objects.filter(
            is_active_classification=True,
            transaction_type__name__in=income_types,
            value_date__range=(m_start, m_end),
            bank_transaction__is_deleted=False,
            bank_transaction__credit_amount__gt=0
        ).select_related('transaction_type', 'cost_centre')
        
        if getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
            bank_income_qs = bank_income_qs.filter(transaction_type__company__in=companies_rel.all())

        for item in bank_income_qs:
            amount = item.amount or Decimal('0')
            type_name = item.transaction_type.name
            cc_name = item.cost_centre.name if item.cost_centre else "Other"
            
            igen_income_type_breakdown[type_name] += amount
            igen_income_cc_breakdown[cc_name] += amount

        # 2. CASH Income Credits
        cash_income_filters = Q(
            is_active=True,
            transaction_type__name__in=income_types,
            date__range=(m_start, m_end),
            amount__gt=0
        )
        if getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
            cash_income_filters &= Q(company__in=companies_rel.all())

        cash_income_qs = CashLedgerRegister.objects.filter(cash_income_filters).select_related('transaction_type', 'cost_centre')
        for item in cash_income_qs:
            amount = item.amount or Decimal('0')
            type_name = item.transaction_type.name
            cc_name = item.cost_centre.name if item.cost_centre else "Other"
            
            igen_income_type_breakdown[type_name] += amount
            igen_income_cc_breakdown[cc_name] += amount

        # Total Calculation
        total_income_credits = sum(igen_income_type_breakdown.values(), Decimal('0'))
        total_igen_income = total_income_credits + total_margin_collected_sum

        # Formatting for response
        formatted_income_type_breakdown = [
            {"type": k, "amount": v} for k, v in igen_income_type_breakdown.items()
        ]
        formatted_income_cc_breakdown = [
            {"cost_centre": k, "amount": v} for k, v in igen_income_cc_breakdown.items()
        ]

        # Total iGen Expenses (9th Enhancement)
        # Filters: debit only, excluding Rental/Sale/Maintenance cost centres
        expense_regex = r'(igen service charge|cleaning|stationery|transport|fuel|office rent)'
        igen_expense_type_breakdown = defaultdict(Decimal)
        igen_expense_cc_breakdown = defaultdict(Decimal)

        # 1. BANK Expenses (Debits)
        bank_expense_qs = Classification.objects.filter(
            is_active_classification=True,
            transaction_type__name__iregex=expense_regex,
            value_date__range=(m_start, m_end),
            bank_transaction__is_deleted=False,
            bank_transaction__debit_amount__gt=0
        ).exclude(
            cost_centre__name__iregex=r'(Rental|Sale|Maintenance)'
        ).select_related('transaction_type', 'cost_centre')

        if getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
            bank_expense_qs = bank_expense_qs.filter(transaction_type__company__in=companies_rel.all())

        for item in bank_expense_qs:
            # For expenses, we use the debit amount (treated as positive for KPI display)
            amount = item.bank_transaction.debit_amount or Decimal('0')
            type_name = item.transaction_type.name
            cc_name = item.cost_centre.name if item.cost_centre else "Administrative"
            
            igen_expense_type_breakdown[type_name] += amount
            igen_expense_cc_breakdown[cc_name] += amount

        # 2. CASH Expenses (Payments)
        cash_expense_filters = Q(
            is_active=True,
            transaction_type__name__iregex=expense_regex,
            date__range=(m_start, m_end),
            amount__lt=0 # Debits are negative in Cash ledger
        )
        if getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
            cash_expense_filters &= Q(company__in=companies_rel.all())

        cash_expense_qs = CashLedgerRegister.objects.filter(cash_expense_filters).exclude(
            cost_centre__name__iregex=r'(Rental|Sale|Maintenance)'
        ).select_related('transaction_type', 'cost_centre')

        for item in cash_expense_qs:
            amount = abs(item.amount or Decimal('0'))
            type_name = item.transaction_type.name
            cc_name = item.cost_centre.name if item.cost_centre else "Administrative"
            
            igen_expense_type_breakdown[type_name] += amount
            igen_expense_cc_breakdown[cc_name] += amount

        total_igen_expenses = sum(igen_expense_type_breakdown.values(), Decimal('0'))

        payload = {
            "total_properties": total,
            "rented": rented_contribution_count,
            "vacant": total - rented_contribution_count,
            "care": care,
            "sale": sale,
            "rent_to_be_collected": rent_sum.quantize(Decimal("1")), # Theoretical
            "rent_received": rent_received_sum.quantize(Decimal("1")),
            "rent_pending_collection": (rent_sum - rent_received_sum).quantize(Decimal("1")), 
            "igen_sc_this_month": sc_sum.quantize(Decimal("1")),
            "igen_sc_collected": sc_collected_sum.quantize(Decimal("1")),
            "igen_sc_variance": (sc_collected_sum - sc_sum).quantize(Decimal("1")),
            "owner_recoverables_total": owner_recoverables_sum.quantize(Decimal("1")),
            "owner_recoverables_base": owner_recoverables_base.quantize(Decimal("1")),
            "owner_recoverables_margin": owner_recoverables_margin.quantize(Decimal("1")),
            "total_margin_collected": total_margin_collected_sum.quantize(Decimal("1")),
            "total_igen_income": total_igen_income.quantize(Decimal("1")),
            "total_igen_expenses": total_igen_expenses.quantize(Decimal("1")),
            "igen_income_type_breakdown": formatted_income_type_breakdown,
            "igen_income_cc_breakdown": formatted_income_cc_breakdown,
            "igen_expense_type_breakdown": [{"type": k, "amount": v} for k, v in igen_expense_type_breakdown.items()],
            "igen_expense_cc_breakdown": [{"cost_centre": k, "amount": v} for k, v in igen_expense_cc_breakdown.items()],
            "inspections_30d": inspections,
            "inspections_due_5d": inspections_due,
            "inspections_expired": inspections_expired,
            "renewals_30d": renewals_30d,
            "agreements_expired": agreements_expired,
            "to_be_vacated_30d": to_vacate,
            "margin_breakdown": formatted_margin_breakdown,
        }
        return Response(OwnerRentalSummarySerializer(payload).data)


class OwnerRentalPropertiesView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        PropertyModel = _get_property_model()
        if PropertyModel is None:
            return Response([])

        qs = PropertyModel.objects.filter(is_active=True)
        user = request.user
        companies_rel = getattr(user, "companies", None)
        if getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
            try:
                qs = qs.filter(company__in=companies_rel.all())
            except Exception:
                pass

        month_str = request.GET.get("month") or date.today().strftime("%Y-%m")
        m_start, m_end = _month_range(month_str)
        days_in_month = Decimal(str((m_end - m_start).days + 1))

        rows = []
        for p in qs:
            base_rent = getattr(p, "monthly_rent", Decimal("0")) or Decimal("0")
            base_sc = getattr(p, "igen_service_charge", Decimal("0")) or Decimal("0")
            l_start = getattr(p, "lease_start_date", None)
            l_end = getattr(p, "lease_end_date", None)

            status_str = (getattr(p, "status", "") or "").strip().lower()
            
            display_rent = Decimal("0")
            display_sc = Decimal("0")
            
            if status_str == "occupied":
                actual_start = max(l_start, m_start) if l_start else m_start
                actual_end = min(l_end or m_end, m_end)
                if actual_end >= actual_start:
                    occ_days = (actual_end - actual_start).days + 1
                    display_rent = (base_rent / days_in_month) * Decimal(str(occ_days))
                    display_sc = (base_sc / days_in_month) * Decimal(str(occ_days))

            tenant_contact = getattr(p, "tenant_contact", None)
            tenant_name = getattr(tenant_contact, "full_name", None) or getattr(p, "tenant", None)
            owner_name = getattr(getattr(p, "landlord", None), "full_name", None)
            
            # Safe flag retrieval
            flags_obj = getattr(p, "owner_flags", None)
            is_txn_scheduled = getattr(flags_obj, "transaction_scheduled", False) if flags_obj else False
            is_email_sent = getattr(flags_obj, "email_sent", False) if flags_obj else False

            rows.append({
                "id": p.id,
                "property_name": p.name,
                "status": p.status,
                "base_rent": base_rent.quantize(Decimal("1")), # Full master rent
                "base_igen_service_charge": base_sc.quantize(Decimal("1")),
                "rent": display_rent.quantize(Decimal("1")), # Pro-rated display rent
                "igen_service_charge": display_sc.quantize(Decimal("1")), 
                "lease_start": l_start,
                "lease_expiry": l_end,
                "agreement_renewal_date": l_end,
                "inspection_date": getattr(p, "next_inspection_date", None),
                "tenant_or_owner": tenant_name or owner_name,
                "transaction_scheduled": is_txn_scheduled,
                "email_sent": is_email_sent,
                "entity_id": _extract_entity_id_from_property(p),
            })
        return Response(OwnerRentalRowSerializer(rows, many=True).data)


class OwnerRentalPendingPropertiesView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        PropertyModel = _get_property_model()
        if PropertyModel is None:
            return Response({"rows": [], "unmapped_received": "0"})

        # Filters
        user = request.user
        month_str = request.GET.get("month") or date.today().strftime("%Y-%m")
        m_start, m_end = _month_range(month_str)
        days_in_month = Decimal(str((m_end - m_start).days + 1))
        
        qs = PropertyModel.objects.filter(is_active=True)
        companies_rel = getattr(user, "companies", None)
        if getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
             qs = qs.filter(company__in=companies_rel.all())

        # 1. Total Rent Received from all BANK entries (for reconciliation)
        rent_received_filters = Q(
            is_active_classification=True,
            transaction_type__name__iexact='Rent In', # Requirement: Only take 'Rent In'
            value_date__range=(m_start, m_end),
            bank_transaction__source='BANK',
            bank_transaction__is_deleted=False,
            bank_transaction__credit_amount__gt=0
        )
        if getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
            rent_received_filters &= Q(transaction_type__company__in=companies_rel.all())
        
        # 2. Map receipts to properties
        # This gives us property_id -> total received
        receipts = Classification.objects.filter(rent_received_filters).values(
            'entity__linked_property_id'
        ).annotate(total=Sum('amount'))
        
        received_map = {
            r['entity__linked_property_id']: r['total'] 
            for r in receipts if r['entity__linked_property_id']
        }
        
        mapped_total = sum(received_map.values()) if received_map else Decimal("0")
        total_received = Classification.objects.filter(rent_received_filters).aggregate(Sum('amount'))['amount__sum'] or Decimal("0")
        unmapped_total = total_received - mapped_total

        # 3. Process Properties
        pending_list = []
        # Filter for Occupied only (TC-05)
        for p in qs.filter(status__iexact='Occupied'):
            base_rent = getattr(p, "monthly_rent", Decimal("0")) or Decimal("0")
            l_start = getattr(p, "lease_start_date", None)
            l_end = getattr(p, "lease_end_date", None)
            
            # Simple pro-rating
            expected_prop = Decimal("0")
            actual_start = max(l_start, m_start) if l_start else m_start
            actual_end = min(l_end or m_end, m_end)
            if actual_end >= actual_start:
                occ_days = (actual_end - actual_start).days + 1
                expected_prop = (base_rent / days_in_month) * Decimal(str(occ_days))
            
            received_prop = received_map.get(p.id, Decimal("0"))
            pending_prop = expected_prop - received_prop
            
            # Include any property with a non-zero discrepancy (Debt or Advance)
            if pending_prop.quantize(Decimal("1")) != 0:
                tenant_contact = getattr(p, "tenant_contact", None)
                tenant_name = getattr(tenant_contact, "full_name", None) or getattr(p, "tenant", "N/A")
                
                pending_list.append({
                    "property_id": p.id,
                    "property_name": p.name,
                    "tenant_name": tenant_name,
                    "monthly_rent": base_rent.quantize(Decimal("1")),
                    "expected_rent": expected_prop.quantize(Decimal("1")),
                    "received_rent": received_prop.quantize(Decimal("1")),
                    "pending_amount": pending_prop.quantize(Decimal("1")),
                })

        # Sort: Pending (desc), then Name
        pending_list.sort(key=lambda x: (-x['pending_amount'], x['property_name']))

        return Response({
            "rows": OwnerRentalPendingPropertySerializer(pending_list, many=True).data,
            "unmapped_received": unmapped_total.quantize(Decimal("1"))
        })


class OwnerRentalInspectionExpiriesView(APIView):
    """
    Inspection Drill-down for Owner Dashboard:
      - Due list: inspection_due_date within next 5 days
      - Expiry list: inspection_expiry_date within next 5 days
    """
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        PropertyModel = _get_property_model()
        if PropertyModel is None:
            return Response({"rows": []})

        user = request.user
        type_filter = request.GET.get("type", "upcoming")
        company_id = request.GET.get("company_id")
        today = date.today()
        in_5 = today + timedelta(days=5)

        qs = PropertyModel.objects.filter(is_active=True)
        if company_id:
            qs = qs.filter(company_id=company_id)
        else:
            companies_rel = getattr(user, "companies", None)
            if getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
                qs = qs.filter(company__in=companies_rel.all())

        results = []
        if type_filter in ("upcoming", "due"):
            insp_qs = qs.filter(next_inspection_date__gte=today, next_inspection_date__lte=in_5)
        else:
            insp_qs = qs.filter(next_inspection_date__lt=today)
            
        for p in insp_qs:
            days_left = (p.next_inspection_date - today).days
            results.append(self._serialize_prop(p, days_left))
        
        results.sort(key=lambda x: x['inspection_date'] if x['inspection_date'] else today)

        return Response({
            "rows": OwnerRentalInspectionExpiryPropertySerializer(results, many=True).data
        })

    def _serialize_prop(self, p, days_left):
        tenant_contact = getattr(p, "tenant_contact", None)
        tenant_name = getattr(tenant_contact, "full_name", None) or getattr(p, "tenant", "N/A")
        landlord = getattr(p, "landlord", None)
        owner_name = getattr(landlord, "full_name", "N/A")
        pm = getattr(p, "project_manager", None)
        pm_name = getattr(pm, "full_name", "N/A")

        return {
            "property_id": p.id,
            "property_name": p.name,
            "inspection_date": p.next_inspection_date,
            "days_left": days_left,
            "tenant_name": tenant_name,
            "owner_name": owner_name,
            "project_manager": pm_name
        }


class OwnerRentalAgreementExpiriesView(APIView):
    """
    Agreement Renewal/Expiry Drill-down for Owner Dashboard.
    """
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        PropertyModel = _get_property_model()
        if PropertyModel is None:
            return Response({"rows": []})

        user = request.user
        type_filter = (request.GET.get("type") or "upcoming").strip().lower()
        company_id = request.GET.get("company_id")
        today = date.today()
        in_30 = today + timedelta(days=30)

        # Using icontains to handle potential trailing spaces in DB
        qs = PropertyModel.objects.filter(is_active=True, status__icontains="occupied")
        if company_id:
            qs = qs.filter(company_id=company_id)
        else:
            companies_rel = getattr(user, "companies", None)
            if getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
                qs = qs.filter(company__in=companies_rel.all())

        results = []
        if type_filter in ["upcoming", "due", "renewal"]:
            agr_qs = qs.filter(lease_end_date__gte=today, lease_end_date__lte=in_30)
        else:
            agr_qs = qs.filter(lease_end_date__lt=today)
            
        for p in agr_qs:
            days_left = (p.lease_end_date - today).days if p.lease_end_date else 0
            
            tenant_contact = getattr(p, "tenant_contact", None)
            tenant_name = getattr(tenant_contact, "full_name", None) or getattr(p, "tenant", "N/A")
            landlord = getattr(p, "landlord", None)
            owner_name = getattr(landlord, "full_name", "N/A")

            results.append({
                "property_id": p.id,
                "property_name": p.name,
                "expiry_date": p.lease_end_date,
                "days_left": days_left,
                "tenant_name": tenant_name,
                "owner_name": owner_name
            })
        
        # Sorting
        if type_filter == "upcoming":
            results.sort(key=lambda x: x['expiry_date'] if x['expiry_date'] else today)
        else:
            # Overdue - most overdue first (smallest date first)
            results.sort(key=lambda x: x['expiry_date'] if x['expiry_date'] else today)

        return Response({
            "rows": OwnerRentalAgreementExpiryPropertySerializer(results, many=True).data
        })


class OwnerRentalServiceChargeBreakdownView(APIView):
    """
    Service Charge Breakdown Drill-down for Owner Dashboard.
    Groups collected service charges by their specific transaction type names 
    (e.g., Owner vs Tenant) and lists the properties involved.
    """
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        user = request.user
        PropertyModel = _get_property_model()
        selected_month_str = request.GET.get("month")
        company_id = request.GET.get("company_id")

        if not selected_month_str or PropertyModel is None:
            return Response({"rows": [], "summaries": [], "unmapped_total": 0})

        try:
            m_start, m_end = _month_range(selected_month_str)
            days_in_month = Decimal(str((m_end - m_start).days + 1))
        except Exception:
            return Response({"rows": [], "summaries": [], "unmapped_total": 0})

        # 1. Base Filters for Properties
        prop_qs = PropertyModel.objects.filter(is_active=True)
        companies_rel = getattr(user, "companies", None)
        if company_id:
            prop_qs = prop_qs.filter(company_id=company_id)
        elif getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
            prop_qs = prop_qs.filter(company__in=companies_rel.all())

        # 2. Get Collected Service Charge (splits)
        sc_filters = Q(
            is_active_classification=True,
            transaction_type__name__icontains='service charge',
            value_date__range=(m_start, m_end),
            bank_transaction__source='BANK',
            bank_transaction__is_deleted=False,
            bank_transaction__credit_amount__gt=0
        )
        if company_id:
            sc_filters &= Q(transaction_type__company_id=company_id)
        elif getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
            sc_filters &= Q(transaction_type__company__in=companies_rel.all())

        collected_txns = Classification.objects.filter(sc_filters).select_related(
            'transaction_type', 'entity', 'entity__linked_property'
        )

        # 3. Process Collections
        property_collected = defaultdict(lambda: {"total": Decimal("0"), "txns": []})
        summaries_dict = defaultdict(Decimal)
        unmapped_total = Decimal("0")

        for t in collected_txns:
            amt = t.amount or Decimal("0")
            type_name = getattr(t.transaction_type, "name", "Other SC")
            summaries_dict[type_name] += amt

            prop_id = getattr(t.entity, "linked_property_id", None)
            if prop_id:
                property_collected[prop_id]["total"] += amt
                property_collected[prop_id]["txns"].append({
                    "date": t.value_date,
                    "amount": amt,
                    "type": type_name
                })
            else:
                unmapped_total += amt

        # 4. Process Properties for Expected vs Collected
        rows = []
        for p in prop_qs:
            # Pro-rated Expected Logic
            base_sc = getattr(p, "igen_service_charge", Decimal("0")) or Decimal("0")
            l_start = getattr(p, "lease_start_date", None)
            l_end = getattr(p, "lease_end_date", None)

            if l_start and l_end and l_end < l_start:
                continue

            actual_start = max(l_start, m_start) if l_start else m_start
            actual_end = min(l_end or m_end, m_end)
            
            expected = Decimal("0")
            if actual_end >= actual_start:
                occupied_days = (actual_end - actual_start).days + 1
                expected = (base_sc / days_in_month) * Decimal(str(occupied_days))

            collected_data = property_collected.get(p.id, {"total": Decimal("0"), "txns": []})
            collected_amt = collected_data["total"]

            # Only include row if there is something to show
            if expected > 0 or collected_amt > 0:
                tenant_contact = getattr(p, "tenant_contact", None)
                tenant_name = getattr(tenant_contact, "full_name", None) or getattr(p, "tenant", "N/A")
                
                rows.append({
                    "property_id": p.id,
                    "property_name": p.name,
                    "tenant_name": tenant_name,
                    "expected_amount": expected,
                    "collected_amount": collected_amt,
                    "variance": collected_amt - expected,
                    "details": collected_data["txns"]
                })

        # 5. Format Result
        summaries = [{"type_name": k, "total_amount": v} for k, v in summaries_dict.items()]
        summaries.sort(key=lambda x: x['total_amount'], reverse=True)
        
        # Sort rows by largest absolute variance (show discrepancies first)
        rows.sort(key=lambda x: abs(x['variance']), reverse=True)

        return Response({
            "rows": OwnerRentalServiceChargeBreakdownSerializer(rows, many=True).data,
            "summaries": summaries,
            "unmapped_total": unmapped_total
        })


class OwnerRentalMaintenanceBreakdownView(APIView):
    """
    Maintenance/Expenses Breakdown Drill-down for Owner Dashboard.
    Lists all Rental/Sale cost centre entries (BANK + CASH).
    Includes base amount, margin, and total collectible.
    """
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        user = request.user
        selected_month_str = request.GET.get("month")
        company_id = request.GET.get("company_id")

        if not selected_month_str:
            return Response({"rows": []})

        try:
            m_start, m_end = _month_range(selected_month_str)
        except Exception:
            return Response({"rows": []})

        companies_rel = getattr(user, "companies", None)
        
        # 1. BANK Recoverables
        sc_filters = Q(
            is_active_classification=True,
            transaction_type__name__iregex=r'(Maintenance|Legal|Paper Work|Paperwork)',
            value_date__range=(m_start, m_end),
            bank_transaction__is_deleted=False
        )
        if company_id:
            sc_filters &= Q(transaction_type__company_id=company_id)
        elif getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
            sc_filters &= Q(transaction_type__company__in=companies_rel.all())

        bank_txns = Classification.objects.filter(sc_filters).select_related(
            'transaction_type', 'entity', 'entity__linked_property', 'cost_centre'
        )

        # 2. CASH Recoverables
        from cash_ledger.models import CashLedgerRegister
        cash_filters = Q(
            is_active=True,
            transaction_type__name__iregex=r'(Maintenance|Legal|Paper Work|Paperwork)',
            date__range=(m_start, m_end)
        )
        if company_id:
            cash_filters &= Q(company_id=company_id)
        elif getattr(user, "role", None) != "SUPER_USER" and companies_rel is not None:
            cash_filters &= Q(company__in=companies_rel.all())

        cash_txns = CashLedgerRegister.objects.filter(cash_filters).select_related(
            'cost_centre', 'entity', 'entity__linked_property'
        )

        rows = []

        # Process BANK
        for t in bank_txns:
            base = t.amount or Decimal('0')
            mgn = t.parsed_margin or Decimal('0')
            prop = getattr(t.entity, "linked_property", None)
            
            rows.append({
                "property_id": prop.id if prop else None,
                "property_name": prop.name if prop else "Not Linked",
                "cost_centre": t.cost_centre.name if t.cost_centre else "N/A",
                "txn_type": t.transaction_type.name if t.transaction_type else "Bank Trxn",
                "base_amount": base,
                "margin_amount": mgn,
                "total_collectible": base + mgn,
                "date": t.value_date,
                "remarks": t.remarks,
                "source": "BANK"
            })

        # Process CASH
        for c in cash_txns:
            base = c.amount or Decimal('0')
            mgn = c.margin or Decimal('0')
            prop = getattr(c.entity, "linked_property", None)
            
            rows.append({
                "property_id": prop.id if prop else None,
                "property_name": prop.name if prop else "Not Linked",
                "cost_centre": c.cost_centre.name if c.cost_centre else "N/A",
                "txn_type": c.remarks or "Cash Payment",
                "base_amount": base,
                "margin_amount": mgn,
                "total_collectible": base + mgn,
                "date": c.date,
                "remarks": c.remarks,
                "source": "CASH"
            })

        # Sort by date DESC
        rows.sort(key=lambda x: x['date'], reverse=True)

        return Response({
            "rows": OwnerRentalMaintenanceBreakdownSerializer(rows, many=True).data
        })


class OwnerRentalPropertyPatchView(APIView):
    """
    Inline updates for Owner Dashboard:
      - Rent
      - iGen Service Charge
      - Lease Start / Lease Expiry
      - Agreemion_scheduled, email_sent
    """

    permission_classes = [permissions.IsAuthenticated]

    def patch(self, request, pk):
        from django.shortcuts import get_object_or_404
        from rest_framework import status
        from rest_framework.response import Response as DRFResponse

        prop = get_object_or_404(Property, pk=pk)
        data = request.data or {}
        flags, _ = OwnerRentalFlag.objects.get_or_create(property=prop)

        prop_fields_changed = []
        flag_fields_changed = {}

        # ---------- helpers ----------
        def to_bool(v):
            if isinstance(v, bool):
                return v
            if isinstance(v, str):
                return v.strip().lower() in ("1", "true", "t", "yes", "y", "on")
            try:
                return bool(int(v))
            except Exception:
                return False

        def clean_number(val):
            if val in ("", None):
                return None
            if isinstance(val, (int, float, Decimal)):
                return val
            try:
                # allow "25,000" style
                return Decimal(str(val).replace(",", ""))
            except Exception:
                return None

        def apply_rent():
            if "rent" not in data:
                return
            val = clean_number(data.get("rent"))
            # Prefer monthly_rent, then expected_rent, finally rent
            for field in ("monthly_rent", "expected_rent", "rent"):
                if hasattr(prop, field):
                    if getattr(prop, field) != val:
                        setattr(prop, field, val)
                        prop_fields_changed.append(field)
                    break

        def apply_igen_sc():
            if "igen_service_charge" not in data or not hasattr(
                prop,
                "igen_service_charge",
            ):
                return
            val = clean_number(data.get("igen_service_charge"))
            if getattr(prop, "igen_service_charge") != val:
                prop.igen_service_charge = val
                prop_fields_changed.append("igen_service_charge")

        def apply_date(payload_key, candidates):
            if payload_key not in data:
                return
            raw = data.get(payload_key)
            dt = parse_date(raw) if raw else None
            for field in candidates:
                if hasattr(prop, field):
                    if getattr(prop, field) != dt:
                        setattr(prop, field, dt)
                        prop_fields_changed.append(field)
                    break

        # ---------- apply property edits ----------
        apply_rent()
        apply_igen_sc()
        apply_date("lease_start", ["lease_start_date", "lease_start"])
        apply_date("lease_expiry", ["lease_end_date", "lease_expiry"])
        apply_date("agreement_renewal_date", ["agreement_renewal_date"])

        # ---------- flags ----------
        if "transaction_scheduled" in data or "txn_scheduled" in data:
            v = data.get("transaction_scheduled", data.get("txn_scheduled"))
            nv = to_bool(v)
            if flags.transaction_scheduled != nv:
                flags.transaction_scheduled = nv
                flag_fields_changed["transaction_scheduled"] = nv

        if "email_sent" in data:
            nv = to_bool(data["email_sent"])
            if flags.email_sent != nv:
                flags.email_sent = nv
                flag_fields_changed["email_sent"] = nv

        # ---------- save ----------
        if prop_fields_changed:
            prop.save(update_fields=list(set(prop_fields_changed)))
        if flag_fields_changed:
            flags.save(update_fields=list(flag_fields_changed.keys()))

        return DRFResponse(
            {
                "status": "ok",
                "property_fields": list(set(prop_fields_changed)),
                "flags": flag_fields_changed,
            },
            status.HTTP_200_OK,
        )


# -------- NEW: Property-based Statement (used by Owner Dashboard "Generate") --------


class OwnerRentalPropertyStatementPDFView(APIView):
    """
    Generate statement for a single property as PDF.
    Params (any one period selector):
      - property_id (required)
      - month (YYYY-MM)  OR
      - from (YYYY-MM-DD) & to (YYYY-MM-DD)
    """

    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        property_id = request.GET.get("property_id")
        month = request.GET.get("month")
        from_str = request.GET.get("from")
        to_str = request.GET.get("to")

        if not property_id:
            return Response(
                {"detail": "property_id is required"},
                status=400,
            )

        try:
            prop = Property.objects.get(pk=int(property_id))
        except (Property.DoesNotExist, ValueError):
            return Response({"detail": "Property not found"}, status=404)

        # Determine period: explicit from/to win over month
        from_date = to_date = None
        period_label = None

        if from_str or to_str:
            from_date = parse_date(from_str) if from_str else None
            to_date = parse_date(to_str) if to_str else None
            if not (from_date and to_date):
                return Response(
                    {"detail": "Both 'from' and 'to' must be valid dates (YYYY-MM-DD)."},
                    status=400,
                )
            period_label = f"{from_date} to {to_date}"
            fallback_month = month or from_date.strftime("%Y-%m")
        else:
            if not month:
                return Response(
                    {"detail": "Either month (YYYY-MM) or from/to is required."},
                    status=400,
                )
            from_date, to_date = _statement_5th_range(month)
            period_label = f"{from_date} to {to_date}"
            fallback_month = month

        rows = _statement_rows_for_property(request.user, prop, from_date, to_date)

        if not rows:
           rows = [
                {
                    "value_date": from_date,
                    "txn_type": "",
                    "credit": Decimal("0"),
                    "debit": Decimal("0"),
                    "balance": Decimal("0"),
                    "remarks": f"No transactions recorded for {from_date} to {to_date}.",
                }
            ]


        # --- HIDE ONLY tenant-side flows in Owner Statement PDF ---
        # 1) "iGen service charge from tenant"
        # 2) "Token received from tenant"
        def _skip_owner_row(r):
            t = (r.get("txn_type") or "").strip().lower()
            rem = (r.get("remarks") or "").strip().lower()

            # exact tenant-side service charge
            if "igen service charge from tenant" in t or "igen service charge from tenant" in rem:
                return True

            # tenant-side token received
            if "token received from tenant" in t or "token received from tenant" in rem:
                return True

            return False

        rows = [r for r in rows if not _skip_owner_row(r)]

        headers = ["Date", "Type", "Credit", "Debit", "Balance", "Remarks"]
        table = [
            [
                r["value_date"].strftime("%d/%m/%Y") if r.get("value_date") else "",
                r.get("txn_type") or "",
                str(r.get("credit") or 0),
                str(r.get("debit") or 0),
                str(r.get("balance") or 0),
                (r.get("remarks") or "")[:64],
            ]
            for r in rows
        ]

        # Owner + property code for heading and filename
        owner_name = (
            getattr(getattr(prop, "landlord", None), "full_name", None)
            or getattr(getattr(prop, "landlord", None), "name", None)
        )
        prop_code = (
            getattr(prop, "code", None)
            or getattr(prop, "property_code", None)
            or getattr(prop, "name", None)
            or f"Property {prop.id}"
        )
        owner_or_entity = owner_name or "Owner"

        # Heading: "OwnerName - PropertyCode - Period Owner Statement"
        title = f"{owner_or_entity} - {prop_code} Owner Statement"

        pdf = export_simple_pdf(title, headers, table)
        resp = HttpResponse(pdf.read(), content_type="application/pdf")

        if month:
            suffix = month
        else:
            suffix = f"{from_date}_to_{to_date}"

        fname = f"OwnerStatement_{_safe_filename_part(prop_code)}_{suffix}.pdf"
        resp["Content-Disposition"] = f'attachment; filename="{fname}"'
        return resp




class OwnerRentalPropertyStatementDOCXView(APIView):
    """
    Generate statement for a single property as Word.
    Params (any one period selector):
      - property_id (required)
      - month (YYYY-MM)  OR
      - from (YYYY-MM-DD) & to (YYYY-MM-DD)
    """

    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        if not _DOCX_OK:
            return Response(
                {"detail": "Word export unavailable: install python-docx"},
                status=500,
            )

        property_id = request.GET.get("property_id")
        month = request.GET.get("month")
        from_str = request.GET.get("from")
        to_str = request.GET.get("to")

        if not property_id:
            return Response(
                {"detail": "property_id is required"},
                status=400,
            )

        try:
            prop = Property.objects.get(pk=int(property_id))
        except (Property.DoesNotExist, ValueError):
            return Response({"detail": "Property not found"}, status=404)

        # Determine period
        from_date = to_date = None
        period_label = None

        if from_str or to_str:
            from_date = parse_date(from_str) if from_str else None
            to_date = parse_date(to_str) if to_str else None
            if not (from_date and to_date):
                return Response(
                    {"detail": "Both 'from' and 'to' must be valid dates (YYYY-MM-DD)."},
                    status=400,
                )
            period_label = f"{from_date} to {to_date}"
            fallback_month = month or from_date.strftime("%Y-%m")
        else:
            if not month:
                return Response(
                    {"detail": "Either month (YYYY-MM) or from/to is required."},
                    status=400,
                )
            from_date, to_date = _statement_5th_range(month)
            period_label = f"{from_date} to {to_date}"
            fallback_month = month

        rows = _statement_rows_for_property(request.user, prop, from_date, to_date)
        if not rows and fallback_month:
            rows = _synthetic_property_statement_rows(prop, fallback_month)

        # --- HIDE ONLY tenant-side flows in Owner Statement DOCX ---
        # 1) "iGen service charge from tenant"
        # 2) "Token received from tenant"
        def _skip_owner_row(r):
            t = (r.get("txn_type") or "").strip().lower()
            rem = (r.get("remarks") or "").strip().lower()

            if "igen service charge from tenant" in t or "igen service charge from tenant" in rem:
                return True

            if "token received from tenant" in t or "token received from tenant" in rem:
                return True

            return False

        rows = [r for r in rows if not _skip_owner_row(r)]

        owner_name = (
            getattr(getattr(prop, "landlord", None), "full_name", None)
            or getattr(getattr(prop, "landlord", None), "name", None)
        )
        prop_code = (
            getattr(prop, "code", None)
            or getattr(prop, "property_code", None)
            or getattr(prop, "name", None)
            or f"Property {prop.id}"
        )
        owner_or_entity = owner_name or "Owner"

        doc = Document()
        doc.add_heading(
            f"{owner_or_entity} - {prop_code} Owner Statement", level=1


        )
        p = doc.add_paragraph()
        p.add_run(f"Period: {period_label}").italic = True

        headers = ["Date", "Type", "Credit", "Debit", "Balance", "Remarks"]
        table = doc.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            run = hdr_cells[i].paragraphs[0].add_run(h)
            if Pt:
                run.font.bold = True
                run.font.size = Pt(10)

        for r in rows:
            row_cells = table.add_row().cells
            row_cells[0].text = (
                r["value_date"].strftime("%Y-%m-%d") if r.get("value_date") else ""
            )
            row_cells[1].text = str(r.get("txn_type") or "")
            row_cells[2].text = str(r.get("credit") or 0)
            row_cells[3].text = str(r.get("debit") or 0)
            row_cells[4].text = str(r.get("balance") or 0)
            row_cells[5].text = (r.get("remarks") or "")[:256]

        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        resp = HttpResponse(
            bio.getvalue(),
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )

        if month:
            suffix = month
        else:
            suffix = f"{from_date}_to_{to_date}"

        fname = f"OwnerStatement_{_safe_filename_part(prop_code)}_{suffix}.docx"
        resp["Content-Disposition"] = f'attachment; filename="{fname}"'
        return resp


class OwnerRentalPropertyStatementExcelView(APIView):
    """
    Generate statement for a single property as Excel.
    Params (any one period selector):
      - property_id (required)
      - month (YYYY-MM)  OR
      - from (YYYY-MM-DD) & to (YYYY-MM-DD)
    """

    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        property_id = request.GET.get("property_id")
        month = request.GET.get("month")
        from_str = request.GET.get("from")
        to_str = request.GET.get("to")

        if not property_id:
            return Response(
                {"detail": "property_id is required"},
                status=400,
            )

        try:
            prop = Property.objects.get(pk=int(property_id))
        except (Property.DoesNotExist, ValueError):
            return Response({"detail": "Property not found"}, status=404)

        # Determine period
        from_date = to_date = None
        period_label = None

        if from_str or to_str:
            from_date = parse_date(from_str) if from_str else None
            to_date = parse_date(to_str) if to_str else None
            if not (from_date and to_date):
                return Response(
                    {"detail": "Both 'from' and 'to' must be valid dates (YYYY-MM-DD)."},
                    status=400,
                )
            period_label = f"{from_date} to {to_date}"
            fallback_month = month or from_date.strftime("%Y-%m")
        else:
            if not month:
                return Response(
                    {"detail": "Either month (YYYY-MM) or from/to is required."},
                    status=400,
                )
            from_date, to_date = _statement_5th_range(month)
            period_label = f"{from_date} to {to_date}"
            fallback_month = month

        rows = _statement_rows_for_property(request.user, prop, from_date, to_date)
        if not rows and fallback_month:
            rows = _synthetic_property_statement_rows(prop, fallback_month)

        headers = ["Date", "Type", "Credit", "Debit", "Balance", "Remarks"]
        data = [
            [
                r["value_date"].strftime("%Y-%m-%d") if r.get("value_date") else "",
                r.get("txn_type") or "",
                r.get("credit") or 0,
                r.get("debit") or 0,
                r.get("balance") or 0,
                (r.get("remarks") or "")[:256],
            ]
            for r in rows
        ]

        xlsx = export_excel(headers, data)
        resp = HttpResponse(
            xlsx.read(),
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "spreadsheetml.sheet"
            ),
        )
        prop_code = (
            getattr(prop, "code", None)
            or getattr(prop, "property_code", None)
            or getattr(prop, "name", None)
            or f"Property {prop.id}"
        )

        if month:
            suffix = month
        else:
            suffix = f"{from_date}_to_{to_date}"

        fname = f"OwnerStatement_{_safe_filename_part(prop_code)}_{suffix}.xlsx"
        resp["Content-Disposition"] = f'attachment; filename="{fname}"'
        return resp


# --------------------------- Report 4: Project Profitability ---------------------------


class ProjectProfitabilitySummaryView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        f = parse_date(request.GET.get("from")) or date(date.today().year, 1, 1)
        t = parse_date(request.GET.get("to")) or date.today()
        project_id = request.GET.get("project_id")

        rows = unified_ledger(
            request.user,
            from_date=f,
            to_date=t,
            project_id=int(project_id) if project_id else None,
        )

        # Normalise project name so we always have r["project"]
        for r in rows:
            r["project"] = (
                r.get("project")
                or r.get("project_name")
                or r.get("project__name")
                or "—"
            )

        # group by actual project_id + project name
        agg = {}
        for r in rows:
            key = (r.get("project_id"), r.get("project") or "—")
            a = agg.setdefault(key, {"in": Decimal("0"), "out": Decimal("0")})
            a["in"] += r.get("credit") or Decimal("0")
            a["out"] += r.get("debit") or Decimal("0")

        out = [
            {
                "project_id": k[0],
                "project": k[1],
                "inflows": v["in"],
                "outflows": v["out"],
                "net": v["in"] - v["out"],
            }
            for k, v in agg.items()
        ]
        out.sort(key=lambda x: x["project"] or "")
        return Response(ProjectProfitRowSerializer(out, many=True).data)


class ProjectProfitabilityTransactionsView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        f = parse_date(request.GET.get("from")) or date(date.today().year, 1, 1)
        t = parse_date(request.GET.get("to")) or date.today()
        project_id = request.GET.get("project_id")

        rows = unified_ledger(
            request.user,
            from_date=f,
            to_date=t,
            project_id=int(project_id) if project_id else None,
        )

        # Normalise project name here too so exports & API rows see it
        for r in rows:
            r["project"] = (
                r.get("project")
                or r.get("project_name")
                or r.get("project__name")
                or ""
            )

        rows = running_balance(rows, opening_balance=Decimal("0"))
        return Response(rows)


class ProjectProfitabilityExportView(APIView):
    """
    Excel export for the summary (first) table.
    """

    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        f = parse_date(request.GET.get("from")) or date(date.today().year, 1, 1)
        t = parse_date(request.GET.get("to")) or date.today()
        rows = ProjectProfitabilitySummaryView().get(request).data  # reuse payload

        xlsx = export_excel(
            ["Project ID", "Project", "Inflows", "Outflows", "Net"],
            [
                [
                    r.get("project_id"),
                    r.get("project"),
                    r.get("inflows"),
                    r.get("outflows"),
                    r.get("net"),
                ]
                for r in rows
            ],
        )
        resp = HttpResponse(
            xlsx.read(),
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "spreadsheetml.sheet"
            ),
        )
        resp[
            "Content-Disposition"
        ] = f'attachment; filename="project_profitability_{f}_{t}.xlsx"'
        return resp


class ProjectProfitabilityTransactionsExportView(APIView):
    """
    Excel export for the transactions (third) table.
    Params:
      - from
      - to
      - project_id (optional, filters when provided)
    """

    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        f = parse_date(request.GET.get("from")) or date(date.today().year, 1, 1)
        t = parse_date(request.GET.get("to")) or date.today()
        project_id = request.GET.get("project_id")

        rows = unified_ledger(
            request.user,
            from_date=f,
            to_date=t,
            project_id=int(project_id) if project_id else None,
        )

        # Normalise project name for export
        for r in rows:
            r["project"] = (
                r.get("project")
                or r.get("project_name")
                or r.get("project__name")
                or ""
            )

        rows = running_balance(rows, opening_balance=Decimal("0"))

        headers = [
            "Date",
            "Project",
            "Txn Type",
            "Credit",
            "Debit",
            "Balance",
            "Entity",
            "Cost Centre",
            "Contract",
            "Asset",
            "Remarks",
        ]
        data = []
        for r in rows:
            data.append(
                [
                    r.get("value_date").strftime("%Y-%m-%d") if r.get("value_date") else "",
                    r.get("project") or "",
                    r.get("txn_type") or "",
                    r.get("credit") or 0,
                    r.get("debit") or 0,
                    r.get("balance") or 0,
                    r.get("entity") or "",
                    r.get("cost_centre") or "",
                    r.get("contract") or "",
                    r.get("asset") or "",
                    (r.get("remarks") or "")[:256],
                ]
            )

        xlsx = export_excel(headers, data)
        resp = HttpResponse(
            xlsx.read(),
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "spreadsheetml.sheet"
            ),
        )
        suffix = f"_{f}_{t}"
        if project_id:
            suffix += f"_project{project_id}"
        resp[
            "Content-Disposition"
        ] = f'attachment; filename="project_profitability_transactions{suffix}.xlsx"'
        return resp


# --------------------------- Report 5: Financial Dashboard (Pivot) ---------------------------


class FinancialDashboardPivotView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        body = request.data or []
        # e.g. ["cost_centre","txn_type","entity","asset","contract","date"]
        dims = body.get("dims") or []
        values = body.get("values") or {}  # {"entity":[...], ...}
        date_gran = (
            body.get("date_granularity")
            or body.get("date_gran")
            or body.get("granularity")
            or "day"
        )
        f = parse_date(body.get("from")) or date(date.today().year, 1, 1)
        t = parse_date(body.get("to")) or date.today()

        rows = unified_ledger(request.user, from_date=f, to_date=t)

        # filter
        def allowed(r):
            for d in dims:
                sel = values.get(d)
                if d == "date":
                    val = _format_period(r.get("value_date"), date_gran)
                else:
                    val = _dim_value(r, d) or "—"
                if sel and val not in set(sel):
                    return False
            return True

        filt = [r for r in rows if allowed(r)]

        # group
        grp = defaultdict(lambda: {"credit": Decimal("0"), "debit": Decimal("0")})
        for r in filt:
            key_items = []
            for d in dims:
                if d == "date":
                    key_items.append(_format_period(r.get("value_date"), date_gran))
                else:
                    key_items.append(_dim_value(r, d) or "—")
            key = tuple(key_items)
            grp[key]["credit"] += r.get("credit") or 0
            grp[key]["debit"] += r.get("debit") or 0

        out = []
        for key, v in grp.items():
            item = {dims[i]: key[i] for i in range(len(dims))}
            item["credit"] = v["credit"]
            item["debit"] = v["debit"]
            item["margin"] = v["credit"] - v["debit"]
            out.append(item)

        # running balance across filtered set (by value date)
        bal = Decimal("0")
        for r in sorted(filt, key=lambda x: x["value_date"]):
            bal += (r.get("credit") or 0) - (r.get("debit") or 0)

        totals = {
            "credit": sum([i["credit"] for i in out], Decimal("0")),
            "debit": sum([i["debit"] for i in out], Decimal("0")),
            "margin": sum([i["margin"] for i in out], Decimal("0")),
            "balance": bal,
        }

        return Response({"rows": out, "totals": totals})


class FinancialDashboardExportView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        payload = FinancialDashboardPivotView().post(request).data
        rows = payload.get("rows", [])
        headers = sorted({k for r in rows for k in r.keys()})
        xlsx = export_excel(headers, [[r.get(h, "") for h in headers] for r in rows])
        resp = HttpResponse(
            xlsx.read(),
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "spreadsheetml.sheet"
            ),
        )
        resp[
            "Content-Disposition"
        ] = 'attachment; filename="financial_dashboard_pivot.xlsx"'
        return resp
